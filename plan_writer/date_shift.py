# -*- coding: utf-8 -*-
"""批量日期偏移：正文+表格全覆盖，正确处理短格式范围（5月20日-21日→5月27日-28日）

用法:
  python date_shift.py <docx> --days 7        # 所有日期 +7 天
  python date_shift.py <docx> --days -3       # 所有日期 -3 天
  python date_shift.py <docx> --find --days 7 # 自动发现桌面最新docx

原理：字典映射替代正则，避免短格式范围漏匹配。覆盖正文段落+表格单元格。
"""
import os, re, sys, argparse
from docx import Document
from _utils import find_latest_docx

def build_date_map(shift_days):
    """构建日期映射表：{旧日期文本: 新日期文本}，支持5月6月跨月"""
    date_map = {}

    # 按绝对日偏移：以5月1日为基准（5/1 = day 0）
    DAYS_IN_MAY = 31
    DAYS_IN_JUN = 30

    def abs_day(month, day):
        if month == 5:
            return day - 1
        elif month == 6:
            return DAYS_IN_MAY + (day - 1)
        elif month == 7:
            return DAYS_IN_MAY + DAYS_IN_JUN + (day - 1)
        return -1

    def day_to_str(abs_d):
        if abs_d < DAYS_IN_MAY:
            return '5月{}日'.format(abs_d + 1)
        elif abs_d < DAYS_IN_MAY + DAYS_IN_JUN:
            return '6月{}日'.format(abs_d - DAYS_IN_MAY + 1)
        else:
            return '7月{}日'.format(abs_d - DAYS_IN_MAY - DAYS_IN_JUN + 1)

    for month in (5, 6):
        max_day = DAYS_IN_MAY if month == 5 else DAYS_IN_JUN
        for day in range(1, max_day + 1):
            old_abs = abs_day(month, day)
            new_abs = old_abs + shift_days
            if new_abs < 0:
                continue
            old_text = '{}月{}日'.format(month, day)
            new_text = day_to_str(new_abs)
            date_map[old_text] = new_text

    return date_map


def fix_short_ranges(text, date_map):
    """处理短格式日期范围：5月27日-21日 → 5月27日-28日"""
    # Match: X月YY日[-~至到]ZZ日 where ZZ has no month prefix
    pattern = re.compile(r'([56]\u6708\d{1,2}\u65e5[-~\u81f3\u5230])(\d{1,2})\u65e5')

    def replacer(m):
        prefix = m.group(1)
        short_day = int(m.group(2))
        # Infer original month from prefix: extract month digit
        prefix_month = int(m.group(0)[0])
        old_text = '{}月{}日'.format(prefix_month, short_day)
        if old_text in date_map:
            return prefix + date_map[old_text]
        # Fallback: simple +shift
        total = short_day + (list(date_map.values())[0].count('月') * 0)  # can't infer shift from date_map
        return prefix + '{}日'.format(short_day)  # will be caught by full replace later

    # Better approach: do a two-pass where we first convert short ranges to full format
    def expand_short(m):
        """Expand '5月27日-21日' → '5月27日-5月21日' before applying date_map"""
        pfx = m.group(1)
        short_day = m.group(2)
        month_char = pfx[0]  # '5' or '6'
        return '{}月{}-{}月{}日'.format(month_char, pfx.split('月')[1].rstrip('日-~至到'),
                                          month_char, short_day)

    expanded = re.sub(r'([56])月(\d{1,2})日[-~至到](\d{1,2})日',
                      lambda m: '{}月{}-{}月{}日'.format(m.group(1), m.group(2), m.group(1), m.group(3)),
                      text)
    return expanded


def set_para_text(para, new_text):
    """替换段落全部文本，保留第一个run格式"""
    if para.runs:
        for r in para.runs[1:]:
            r.text = ''
        para.runs[0].text = new_text
    else:
        para.text = new_text


def shift_dates_in_docx(docx_path, shift_days):
    """正文+表格全覆盖日期偏移"""
    date_map = build_date_map(shift_days)
    doc = Document(docx_path)
    total = 0

    # Process paragraphs
    for p in doc.paragraphs:
        old = p.text
        # Step 1: expand short ranges to full format
        new = fix_short_ranges(old, date_map)
        # Step 2: apply date map
        for old_date, new_date in date_map.items():
            new = new.replace(old_date, new_date)
        if new != old:
            set_para_text(p, new)
            total += 1

    # Process table cells
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    old = p.text
                    new = fix_short_ranges(old, date_map)
                    for old_date, new_date in date_map.items():
                        new = new.replace(old_date, new_date)
                    if new != old:
                        set_para_text(p, new)
                        total += 1

    doc.save(docx_path)
    return total




if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    parser = argparse.ArgumentParser(description='Batch date shift in docx (paragraphs + tables)')
    parser.add_argument('docx', nargs='?', help='Path to docx file (optional with --find)')
    parser.add_argument('--find', action='store_true', help='Auto-find latest docx on Desktop')
    parser.add_argument('--days', type=int, required=True, help='Days to shift (+7 for delay, -3 for advance)')
    args = parser.parse_args()

    docx_path = args.docx
    if args.find or not docx_path:
        docx_path = find_latest_docx()
        if not docx_path:
            print('ERROR: No .docx found on Desktop.')
            sys.exit(1)

    if not os.path.exists(docx_path):
        print('ERROR: File not found: {}'.format(docx_path))
        sys.exit(1)

    count = shift_dates_in_docx(docx_path, args.days)
    print('Shifted {} date references by {} days'.format(count, args.days))
    print('Saved: {}'.format(docx_path))
