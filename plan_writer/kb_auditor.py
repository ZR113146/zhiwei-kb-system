#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
kb_auditor — 规范引用审计引擎，两条管线共享。

修改管线：audit_report() → 四问对照表（梳理，非门禁）
新建管线：resolve_for_chapter() → 章节→正确规范条款（精确输出）

架构规则：只 import kb，不 import 其他技能脚本。
"""

import os, re, sys, json, argparse
from docx import Document as _Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'kb_core'))
from kb import KB, normalize_code, extract_code as _ec
from support_guard import partition_by_support_action
import changelog; changelog.record(__file__, sys.argv)

# v6.18: 原文锚点查询 — jieba + 术语词典用于从规范原文提取技术关键词
import jieba
# 修复 B5: 原指向不存在的 ../../pipeline/scripts/kb_term_map.json (被 os.path.exists 静默吞,
# 致 jieba 术语词典从不加载、锚点提取静默退化)。改指 contracts/term_map.json (从 plan_writer/ 上溯一级)。
_TERM_MAP_PATH = os.path.join(os.path.dirname(__file__), '..', 'contracts', 'term_map.json')
_anchor_terms_loaded = False

def _load_jieba_terms():
    global _anchor_terms_loaded
    if _anchor_terms_loaded:
        return
    if os.path.exists(_TERM_MAP_PATH):
        with open(_TERM_MAP_PATH, 'r', encoding='utf-8') as f:
            tm = json.load(f)
        for v in tm.values():
            if isinstance(v, list):
                for t in v:
                    if len(t) >= 2:
                        jieba.add_word(t)
    _anchor_terms_loaded = True

# 通用低区分度词 (排除后保留真正的技术术语)
_ANCHOR_GENERIC = {
    '规定','施工','工程','结构','设计','要求','标准','采用','一般','检查','检验','材料',
    '试验','安装','技术','控制','处理','安全','质量','条件','系统','设备','进行','单位',
    '方法','测量','部位','位置','厚度','高度','深度','长度','宽度','面积','体积','截面',
    '构造','类型','形式','方式','方案','项目','内容','部分','相应','有关','相关','满足',
    '保证','保护','确保','提供','增加','减少','提高','降低','大于','小于','超过','低于',
    '高于','符合','达到','满足','不应','应','宜','可','必须','不得','构件','支撑','节点',
    '梁','板','柱','墙','框架','主控项目','一般项目','允许偏差','检验批','合格','不合格',
    '水平','垂直','横向','纵向','受力','承载力','强度等级','钢筋','混凝土','水泥','钢管',
    '砌体','砂浆','模板','防水','保温',
}

def _extract_anchor_terms(body_text, max_terms=5):
    """从规范原文提取差异化技术术语 (v6.18 原文锚点)。
    与 _build_ai_hint 的区别: 后者从方案段落提取二手措辞, 前者从规范原文提取一手术语。
    策略: 长词优先(≥3字) → 不在通用词表 → 中文 → 去重。长词区分度天然高于短词。"""
    _load_jieba_terms()
    tokens = [t.strip() for t in jieba.lcut(body_text) if len(t.strip()) >= 3]
    terms, seen = [], set()
    for t in tokens:
        if t not in _ANCHOR_GENERIC and t not in seen and re.match(r'^[\u4e00-\u9fff]', t):
            seen.add(t); terms.append(t)
    if len(terms) < max_terms:
        # 长词不够 → 补充 2 字短词
        short_tokens = [t.strip() for t in jieba.lcut(body_text) if len(t.strip()) == 2]
        for t in short_tokens:
            if t not in _ANCHOR_GENERIC and t not in seen and re.match(r'^[\u4e00-\u9fff]', t):
                seen.add(t); terms.append(t)
            if len(terms) >= max_terms:
                break
    return terms[:max_terms]

# ============================================================
#  九章大纲 — 专题 — 规范类别映射
# ============================================================
CHAPTER_MAP = {
    'ch01': {
        'name': '编制依据',
        'expects': ['standards_list'],
    },
    'ch02': {
        'name': '工程概况',
        'expects': ['project_info', 'scope', 'machinery', 'key_points'],
    },
    'ch03': {
        'name': '施工安排',
        'expects': ['org_structure', 'labor_plan', 'zone_flow', 'schedule'],
    },
    'ch04': {
        'name': '施工准备',
        'expects': ['tech_prep', 'test_plan', 'sample_plan', 'material_list'],
    },
    'ch05': {
        'name': '主要施工方法',
        'expects': [
            'earthwork', 'concrete', 'masonry', 'paving',
            'planting', 'pipeline', 'steel', 'demolition',
        ],
        'topic_standards': {
            'earthwork':   ['JGJ79', 'JGJ180', 'GB50007', 'GB50202'],
            'concrete':    ['GB50204', 'GB50666', 'GB50164'],
            'masonry':     ['GB50203', 'JGJ79'],
            'paving':      ['GB50209', 'CJJ1'],
            'planting':    ['CJJ82', 'GB55014', 'CJJT287'],
            'pipeline':    ['GB50268', 'GB50168'],
            'steel':       ['GB50205', 'GB50755'],
            'demolition':  ['JGJ147'],
        },
    },
    'ch06': {
        'name': '质量要求',
        'expects': ['quality_system', 'qc_table', 'acceptance_division', 'defect_prevention'],
        'topic_standards': {
            'quality_system':    ['GB50300'],
            'acceptance':        ['GB50300', 'GB50202', 'GB50203', 'GB50204', 'GB50209'],
            'defect_prevention': ['CJJ82', 'GB50209'],
        },
    },
    'ch07': {
        'name': '安全管理',
        'expects': ['safety_system', 'process_safety', 'emergency', 'special_safety'],
        'topic_standards': {
            'safety_system':    ['JGJ59'],
            'process_safety':   ['JGJ80', 'JGJ180', 'JGJ33'],
            'emergency':        ['GB50720'],
            'electrical':       ['JGJ46'],
            'metro_protection': ['GB50911'],
        },
    },
    'ch08': {
        'name': '文明施工',
        'expects': ['dust_control', 'noise_control', 'waste', 'seasonal'],
        'topic_standards': {
            'noise': ['GB12523'],
            'fire':  ['GB50720'],
        },
    },
    'ch09': {
        'name': '其他要求',
        'expects': ['schedule', 'product_protection', 'cost_reduction'],
    },
}


# ============================================================
#  引擎核心：三阶段引用解析
# ============================================================

# 剥离比较词前缀
_CMP_WORDS = ['不应大于', '不应小于', '不应低于', '不应高于', '不应少于', '不应多于',
              '不得不大于', '不得小于', '不得低于', '不得高于',
              '不大于', '不小于', '不低于', '不高于', '不少于', '不多于',
              '大于', '小于', '低于', '高于', '少于', '多于',
              '≥', '≤', '≧']
def _strip_cmp(s):
    for w in _CMP_WORDS:
        if s.startswith(w):
            return s[len(w):].strip()
    return s.strip()

# 单位归一化——统一到最小单位比对
_UNIT_TO_MM = {
    'mm': 1, 'cm': 10, 'm': 1000, 'km': 1e6,
    'd': 1, '天': 1, 'h': 0.0417, '年': 365,
    'kPa': 1, 'MPa': 1000, 'Pa': 0.001,
    'kN': 1, 'N': 0.001, 'kg': 1, 't': 1000,
    '℃': 1, '°C': 1, '°': 1,
    '%': 1, '㎡': 1, 'm²': 1, 'm³': 1,
}
def _norm_val(s):
    """归一化数值: '20cm'→200, '200mm'→200, '0.5m'→500"""
    m = re.match(r'([\d.]+)\s*([a-zA-Zμ℃°²³㎡]+|天)', s)
    if not m: return None
    num = float(m.group(1))
    unit = m.group(2)
    factor = _UNIT_TO_MM.get(unit, 1)
    return num * factor

# 提取引用模式
_CITE_RE = re.compile(
    r'(?:依据|按|执行|符合|按照)\s*'
    r'((?:GB|JGJ|CJJ|CECS|CJ|DB|TCECS)\s*/?\s*T?\s*\d+[\.-]\d+(?:-\d+)?)'
    r'(?:[^。\n]{0,80}?)'
    r'('
    r'[≥≤≧]\s*\d+\.?\d*\s*(?:kPa|kN|mm|cm|m|MPa|%|d|天|kg/m³|W/\()|'
    r'(?:不[得应])?\s*(?:[小大低高超于少多]于?)\s*\d+\.?\d*\s*(?:kPa|kN|mm|cm|m|MPa|%|[d天]|℃|\u2103|kg/m³)?'
    r')',
    re.DOTALL
)







def _extract_claim_value(claimed):
    val_clean = _strip_cmp(claimed).strip()
    val_match = re.search(r'([\d.]+)', claimed)
    if not val_match:
        return val_clean, None
    return val_clean, val_match.group(1)

def _normalize_ai_hint(ai_hint):
    if isinstance(ai_hint, str):
        return {'context': ai_hint}
    return ai_hint


def _is_project_adapted_context(ai_hint):
    context = (ai_hint or {}).get('context', '')
    return bool(context and re.search(r'(?:本工程|本项目|本方案|保守取|调整为)', context))

def _new_citation_result():
    return {'in_clause': False, 'clause_exists': None, 'detail': '',
            'source': None, 'clause': None, 'suggestion': '', 'needs_ai': False,
            'citation': None, 'clause_type': '', 'version_status': {}, 'audit_status': 'unknown', 'audit_messages': []}


def _load_clause_full(kb, target_code, target_clause):
    if not kb.exists(target_code):
        return None, None
    if hasattr(kb, 'read_clause_full'):
        data = kb.read_clause_full(target_code, target_clause)
        if data:
            return data, data.get('clause_text') or data.get('text') or None
    text = kb.read_clause(target_code, target_clause)
    return None, text


def _attach_clause_audit(result, data):
    if not data:
        return
    result['citation'] = data.get('citation')
    result['clause_type'] = data.get('clause_type', '')
    result['version_status'] = data.get('version_status', {})
    if data.get('citation'):
        result['audit_status'] = data['citation'].get('audit_status', 'unknown')
        result['audit_messages'] = data['citation'].get('audit_messages', [])






def _claim_value_in_clause_text(val_clean, claimed, clause_text):
    val_in = val_clean in clause_text or claimed in clause_text
    if not val_in:
        unit_m = re.search(r'([a-zA-Z\u2103\u00b0]+)$', val_clean)
        if unit_m:
            unit = unit_m.group(1)
            bare_val = val_clean[:unit_m.start()]
            if bare_val in clause_text and unit in clause_text:
                val_in = True
        if not val_in:
            doc_norm = _norm_val(val_clean)
            if doc_norm is not None:
                kb_vals = re.findall(r'[\d.]+\s*(?:mm|cm|m|kPa|MPa|kN|d|天|%|℃|°C|㎡|m²)', clause_text)
                for kv in kb_vals:
                    kb_norm = _norm_val(kv)
                    if kb_norm is not None and abs(doc_norm - kb_norm) < 0.01:
                        val_in = True
                        break
    return val_in







def _value_exists_in_kb(kb, search_q, val_clean, match_fn):
    results = kb.search(search_q, max_results=12) or kb.search(val_clean, max_results=12, vector_weight=0)
    return any(match_fn(search_result.get('text', '')) for search_result in results) if results else False

def _same_standard_matches(kb, code, expected_standards, search_q, match_fn):
    same = []
    targets = [code] + (expected_standards or [])
    for target in dict.fromkeys(targets):
        parts = re.findall(r'\d+', target)
        if len(parts) < 2:
            md = kb.find_md(target)
            if md:
                parts2 = re.findall(r'\d+', os.path.basename(md))
                if len(parts2) >= 2: parts = parts2
        if len(parts) < 2: continue
        must_key = parts[0] + '-' + parts[1]
        results = kb.search(search_q, max_results=8, must=[must_key])
        for search_result in results:
            if match_fn(search_result.get('text', '')):
                same.append((target, search_result)); break
    return same

def _anchor_enhanced_matches(kb, code, clause_text, match_fn):
    matches = []
    anchor_terms = _extract_anchor_terms(clause_text, max_terms=5)
    if len(anchor_terms) >= 2:
        anchor_q = ' '.join(anchor_terms[:4])
        r_anchor = kb.search(anchor_q, max_results=10, vector_weight=0.4)
        nc = normalize_code(code)
        for search_result in r_anchor:
            fn = search_result.get('file', '')
            if (nc in normalize_code(fn) or
                nc in normalize_code(search_result.get('heading', '')) or
                match_fn(search_result.get('text', ''))):
                matches.append((code, search_result))
                break
        if not matches and r_anchor:
            for search_result in r_anchor[:3]:
                if match_fn(search_result.get('text', '')):
                    matches.append(('anchor', search_result))
                    break
    return matches

def _set_citation_suggestion(result, same, val_exists, clause_ref, mode):
    if same:
        target, search_result = same[0]; heading = search_result.get('heading', '')
        if target == 'anchor':
            src_code = _ec(search_result.get('file', ''))
            target = src_code or 'KB'
        clause_match = re.search(r'(\d+\.\d+(?:\.\d+)?)', heading)
        result['source'] = target; result['clause'] = clause_match.group(1) if clause_match else heading[:30]
        if mode == 'value_mismatch':
            result['suggestion'] = f'值不在{clause_ref}，同规范{target} {heading[:40]}'
        else:
            result['suggestion'] = f'条款{clause_ref}不存在，值在{target} {heading[:40]}'
        result['needs_ai'] = True
    elif val_exists:
        if mode == 'value_mismatch':
            result['suggestion'] = '值不在该条款中，同规范搜索未找到'
        else:
            result['suggestion'] = f'条款{clause_ref}不存在，值在KB中但未能定位'
        result['needs_ai'] = True
    else:
        if mode == 'value_mismatch':
            result['suggestion'] = '值不在条款中，KB也未搜到'
        else:
            result['suggestion'] = f'条款{clause_ref}不存在，值也未找到'
        result['needs_ai'] = True
    return result

def _resolve_citation(kb, code, clause_ref, claimed, ai_hint=None):
    """倒金字塔验证引擎。⚠️ SKILL.md 规则5：搜索逻辑不可绕过。

    流程：AI 语义预判 → 搜索验证。搜索不猜测，AI 不搜索。
    禁止在本函数内做关键词提取、语义分类、跨规范推断。
    仅使用 kb.search / kb.read_clause 做纯粹验证。

    ai_hint = {'search_terms': '混凝土 养护 7d',
               'expected_standards': ['GB50666'],
               'context': '原始段落文本'}
    无 AI 时退化为纯值搜索。"""

    # 向后兼容：string → dict
    ai_hint = _normalize_ai_hint(ai_hint)

    r = _new_citation_result()

    val_clean, val = _extract_claim_value(claimed)
    if val is None:
        return r

    if _is_project_adapted_context(ai_hint):
        r['in_clause'] = True
        r['detail'] = 'project-adapted'
        return r

    search_q = (ai_hint or {}).get('search_terms', val_clean)

    # 辅助：纯验证——只检查文本，不做额外搜索
    def _match(ts):
        return val_clean in ts

    # │ Round 1  宽  —  值在 KB 中存在？
    val_exists = _value_exists_in_kb(kb, search_q, val_clean, _match)

    # │ Round 2  窄  —  值在同规范/预期规范中？
    same = _same_standard_matches(kb, code, (ai_hint or {}).get('expected_standards', []), search_q, _match)

    # │ Round 3  精  —  条款原文验证
    clause_data, orig = _load_clause_full(kb, code, clause_ref)
    _attach_clause_audit(r, clause_data)
    r['clause_exists'] = orig is not None

    if orig:
        # 精确匹配：带单位、完整声称、归一化数值
        val_in = _claim_value_in_clause_text(val_clean, claimed, orig)
        if val_in:
            r['in_clause'] = True; r['detail'] = 'matched'; return r
        # v6.18 原文锚点增强: 用规范原文学术语回搜 (补二手措辞的语义差)
        if not same:
            same = _anchor_enhanced_matches(kb, code, orig, _match)
            if same:
                r['detail'] = 'anchor_enhanced'

        _set_citation_suggestion(r, same, val_exists, clause_ref, 'value_mismatch')
        return r

    # 条款不存在
    _set_citation_suggestion(r, same, val_exists, clause_ref, 'missing_clause')
    return r




def _build_ai_hint(para_text, code, claimed, inter):
    """从段落上下文构建 AI 预判提示。
    提取段落中的技术关键词作为 search_terms。context 用全段以支持适配检测。"""
    raw = re.findall(r'[\u4e00-\u9fff]{2,}', para_text)
    stop = set('的了是在和与或中等对为将已可应需其每该各本此及从由按以被把向到于同跟给让叫使令请让')
    terms = [w for w in raw if w not in stop and not w.isdigit()]
    search_terms = f'{_strip_cmp(claimed)} {" ".join(terms[-4:])}'.strip()
    return {
        'search_terms': search_terms,
        'context': para_text,
    }


# ============================================================
#  章节感知：引用属于哪一章？预期在哪一章？
# ============================================================

def _detect_chapter(para_index, all_paras):
    """根据段落位置推断所在章节。all_paras = [(idx, text), ...]"""
    ch = 'ch01'
    for idx, text in all_paras:
        ch_m = re.match(r'^#?\s*([1-9])\s+', text)
        if ch_m and len(text) < 60:
            ch_num = int(ch_m.group(1))
            if idx <= para_index:
                ch = f'ch0{ch_num}' if ch_num < 10 else f'ch{ch_num}'
    return ch


def _check_chapter_fit(code, actual_chapter):
    """检查规范在方案中的引用章节是否合理。
    返回: (ok: bool, expected_chapter: str, reason: str)"""
    nc = normalize_code(code)
    # 在 CHAPTER_MAP 中查找该规范所属的话题
    for ch_key, ch_info in CHAPTER_MAP.items():
        topics = ch_info.get('topic_standards', {})
        for topic, codes in topics.items():
            if nc in codes:
                if ch_key == actual_chapter:
                    return True, ch_key, f'In expected chapter ({ch_info["name"]})'
                else:
                    return False, ch_key, (
                        f'Expected in {ch_info["name"]}({ch_key}), '
                        f'but found in {CHAPTER_MAP.get(actual_chapter, {}).get("name", actual_chapter)}'
                    )
    return None, None, 'No chapter preference for this standard'


# ============================================================
#  输出层 1: audit_report — 修改管线四问对照表
# ============================================================



def _parse_clause_reference(inter):
    clause_match = re.search(r'(?:表\s*(\d+)|第\s*([\d.]+)条|§\s*([\d.]+)|附录\s*([A-Z]))', inter)
    if not clause_match:
        return None, None
    if clause_match.group(1):   clause_ref = f'表{clause_match.group(1)}'
    elif clause_match.group(2): clause_ref = clause_match.group(2)
    elif clause_match.group(3): clause_ref = clause_match.group(3)
    else:                       clause_ref = clause_match.group(4)
    return clause_ref, clause_match.group()



def _retry_citation_with_ai_hint(kb, text, code, clause_ref, claimed, resolution):
    if resolution['in_clause']:
        return resolution
    resolution['needs_ai'] = True
    try:
        from llm_hint import call_llm_for_hint
        llm = call_llm_for_hint(
            text, code, clause_ref, claimed,
            resolution.get('suggestion', ''))
        if llm and 'error' not in llm:
            expected = [s.strip() for s in llm.get('expected_standards', [])]
            ai_hint = {
                'search_terms': llm.get('search_terms', ''),
                'expected_standards': expected,
                'context': text,
            }
            resolved_ai = _resolve_citation(kb, code, clause_ref, claimed, ai_hint)
            if resolved_ai.get('source') or resolved_ai['in_clause']:
                resolution['ai_source'] = resolved_ai.get('source')
                resolution['ai_clause'] = resolved_ai.get('clause')
                resolution['ai_suggestion'] = resolved_ai.get('suggestion')
                resolution['needs_ai'] = False
    except ImportError:
        pass
    except Exception:
        pass
    return resolution


def _render_audit_entry_lines(entry):
    lines = []
    resolution = entry['resolution']
    status = '✓' if resolution['in_clause'] else '✗'
    lines.append(f'\n  [{status}] P{entry["para"]}  {entry["code"]} {entry["clause_display"]}')
    lines.append(f'      方案声称  {entry["claimed"]}')
    if resolution['clause_exists'] is False:
        lines.append(f'      条款存在  ✗ 该条款在 {entry["code"]} 中不存在')
    lines.append(f'      值在条款  {"是" if resolution["in_clause"] else "否"}{" — "+resolution["detail"] if resolution.get("detail") else ""}')
    if resolution.get('clause_type') and resolution.get('clause_type') != 'normative':
        lines.append(f'      条文类型  △ {resolution["clause_type"]}')
    version_status = resolution.get('version_status') or {}
    status_name = version_status.get('status', '') if isinstance(version_status, dict) else ''
    if status_name and status_name != 'effective':
        lines.append(f'      版本状态  △ {status_name}')
    if not resolution['in_clause']:
        if resolution.get('source') and resolution.get('clause'):
            lines.append(f'      同规范中  {resolution["source"]} §{resolution["clause"]}')
        if resolution['suggestion']:
            lines.append(f'      → {resolution["suggestion"]}')
    if entry['chapter_ok'] is False:
        lines.append(f'      章节匹配  △ {entry["chapter_reason"]}')
    return lines


def _append_audit_report_entries(lines, entries):
    current_ch = None
    for entry in entries:
        ch = entry['chapter']
        ch_name = CHAPTER_MAP.get(ch, {}).get('name', ch)
        if ch != current_ch:
            current_ch = ch
            lines.append(f'\n── {ch_name}（{ch}）──')
        lines.extend(_render_audit_entry_lines(entry))


def _append_audit_report_summary(lines, entries):
    total = len(entries)
    ok = sum(1 for entry in entries if entry['resolution']['in_clause'])
    lines.append(f'\n{"=" * 70}')
    lines.append(f'共 {total} 处引用，{ok} 处条款匹配，{total - ok} 处需核查')
    lines.append('=' * 70)

def audit_report(docx_path, kb=None):
    """梳理模式：输出全量四问对照表。返回 (report_lines, issues)"""
    if kb is None:
        kb = KB()
    doc = _Document(docx_path)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    all_paras = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs) if p.text.strip()]

    lines = []
    lines.append('=' * 70)
    lines.append(f'引用审计：{os.path.basename(docx_path)}')
    lines.append('=' * 70)

    entries = []
    for i, text in all_paras:
        for m in _CITE_RE.finditer(text):
            code = m.group(1).strip()
            claimed = m.group(2).strip()
            inter = text[m.start(1):m.start(2)]
            clause_ref, clause_display = _parse_clause_reference(inter)
            if not clause_ref:
                continue

            ch = _detect_chapter(i, all_paras)
            res = _resolve_citation(kb, code, clause_ref, claimed)
            # AI后置接入：引擎失败时调LLM获取语义预判，重跑引擎
            res = _retry_citation_with_ai_hint(kb, text, code, clause_ref, claimed, res)
            ch_ok, ch_exp, ch_reason = _check_chapter_fit(code, ch)

            entry = {
                'para': i, 'code': code, 'clause_ref': clause_ref,
                'clause_display': clause_display, 'claimed': claimed,
                'chapter': ch, 'resolution': res,
                'chapter_ok': ch_ok, 'chapter_expected': ch_exp,
                'chapter_reason': ch_reason,
            }
            entries.append(entry)

    # 按章节分组输出
    _append_audit_report_entries(lines, entries)

    # 统计
    _append_audit_report_summary(lines, entries)

    return lines, entries


def _new_citation_summary(entries):
    return {
        'total': len(entries),
        'matched': 0,
        'needs_review': 0,
        'audit_status': {'pass': 0, 'warn': 0, 'fail': 0, 'unknown': 0},
        'version_status': {},
        'clause_type': {},
        'items': [],
    }


def _citation_review_reasons(in_clause, audit_status, version_name, clause_type, chapter_ok):
    review_reasons = []
    if not in_clause:
        review_reasons.append('value_not_in_clause')
    if audit_status in {'warn', 'fail', 'unknown'}:
        review_reasons.append(f'audit_status={audit_status}')
    if version_name not in {'effective'}:
        review_reasons.append(f'version_status={version_name}')
    if clause_type != 'normative':
        review_reasons.append(f'clause_type={clause_type}')
    if chapter_ok is False:
        review_reasons.append('chapter_mismatch')
    return review_reasons


def _citation_summary_item(entry, citation, resolution, in_clause, audit_status, version_name, clause_type, review_reasons):
    return {
        'para': entry.get('para'),
        'code': entry.get('code', ''),
        'clause_ref': entry.get('clause_ref', ''),
        'claimed': entry.get('claimed', ''),
        'in_clause': in_clause,
        'standard_name': citation.get('standard_name', ''),
        'official_code': citation.get('official_code', ''),
        'clause_type': clause_type,
        'version_status': version_name,
        'audit_status': audit_status,
        'source_file': citation.get('source_file', '') or resolution.get('source', ''),
        'review_reasons': review_reasons,
    }


def citation_audit_summary(entries):
    """Summarize structured citation audit entries for downstream reports."""
    summary = _new_citation_summary(entries)
    for entry in entries:
        resolution = entry.get('resolution', {})
        citation = resolution.get('citation') or {}
        in_clause = bool(resolution.get('in_clause'))
        if in_clause:
            summary['matched'] += 1
        audit_status = citation.get('audit_status') or resolution.get('audit_status') or 'unknown'
        if audit_status not in summary['audit_status']:
            audit_status = 'unknown'
        summary['audit_status'][audit_status] += 1
        version = citation.get('version_status') or resolution.get('version_status') or {}
        version_name = version.get('status', 'unknown') if isinstance(version, dict) else 'unknown'
        summary['version_status'][version_name] = summary['version_status'].get(version_name, 0) + 1
        clause_type = citation.get('clause_type') or resolution.get('clause_type') or 'unknown'
        summary['clause_type'][clause_type] = summary['clause_type'].get(clause_type, 0) + 1
        review_reasons = _citation_review_reasons(
            in_clause, audit_status, version_name, clause_type, entry.get('chapter_ok'))
        if review_reasons:
            summary['needs_review'] += 1
        summary['items'].append(_citation_summary_item(
            entry, citation, resolution, in_clause, audit_status, version_name, clause_type, review_reasons))
    return summary


def _summary_markdown(summary, docx_path):
    lines = [
        '# 引用审计摘要',
        '',
        f'- 文档：{os.path.basename(docx_path)}',
        f'- 引用总数：{summary["total"]}',
        f'- 条款匹配：{summary["matched"]}',
        f'- 需复核：{summary["needs_review"]}',
        f'- 审计状态：pass {summary["audit_status"].get("pass", 0)} / warn {summary["audit_status"].get("warn", 0)} / fail {summary["audit_status"].get("fail", 0)} / unknown {summary["audit_status"].get("unknown", 0)}',
        '',
        '## 需复核引用',
        '',
    ]
    review_items = [item for item in summary['items'] if item.get('review_reasons')]
    if not review_items:
        lines.append('- 无')
    else:
        for item in review_items:
            code = item.get('official_code') or item.get('code')
            reasons = ', '.join(item.get('review_reasons', []))
            lines.append(f"- P{item.get('para')} {code} §{item.get('clause_ref')}：{reasons}")
    lines.extend(['', '## 引用清单', ''])
    for item in summary['items']:
        code = item.get('official_code') or item.get('code')
        lines.append(
            f"- P{item.get('para')} {code} §{item.get('clause_ref')}："
            f"{item.get('audit_status')} / {item.get('version_status')} / {item.get('clause_type')}"
        )
    return '\n'.join(lines) + '\n'


def write_citation_audit_summary(docx_path, entries, output_base):
    """Write citation audit summary as JSON and Markdown. output_base has no suffix."""
    summary = citation_audit_summary(entries)
    os.makedirs(os.path.dirname(os.path.abspath(output_base)), exist_ok=True)
    json_path = output_base + '.json'
    md_path = output_base + '.md'
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(_summary_markdown(summary, docx_path))
    return {'json': json_path, 'md': md_path, 'summary': summary}


def _citation_audit_docx_output_path(docx_path, output_path):
    if output_path is None:
        root, ext = os.path.splitext(docx_path)
        output_path = root + '_citation_audit' + ext
    return output_path


def _add_citation_audit_docx_header(doc, summary, Pt, WD_ALIGN_PARAGRAPH):
    doc.add_page_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('引用审计摘要')
    run.bold = True
    run.font.size = Pt(16)

    metrics = doc.add_paragraph()
    metrics.paragraph_format.first_line_indent = Pt(0)
    metrics.add_run(
        f'引用总数：{summary["total"]}；条款匹配：{summary["matched"]}；'
        f'需复核：{summary["needs_review"]}；'
        f'通过：{summary["audit_status"].get("pass", 0)}；'
        f'警告：{summary["audit_status"].get("warn", 0)}；'
        f'失败：{summary["audit_status"].get("fail", 0)}；'
        f'未知：{summary["audit_status"].get("unknown", 0)}。'
    )


def _add_review_items_docx_table(doc, review_items, Pt, WD_TABLE_ALIGNMENT):
    if review_items:
        h = doc.add_paragraph()
        h.paragraph_format.first_line_indent = Pt(0)
        h.add_run('需复核引用').bold = True
        table = doc.add_table(rows=1, cols=4, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, header in enumerate(['段落', '标准/条文', '声称值', '复核原因']):
            table.rows[0].cells[idx].text = header
        for item in review_items[:30]:
            row = table.add_row().cells
            code = item.get('official_code') or item.get('code')
            row[0].text = f"P{item.get('para')}"
            row[1].text = f"{code} §{item.get('clause_ref')}"
            row[2].text = item.get('claimed', '')
            row[3].text = ', '.join(item.get('review_reasons', []))
    else:
        ok = doc.add_paragraph()
        ok.paragraph_format.first_line_indent = Pt(0)
        ok.add_run('需复核引用：无')


def _add_citation_items_docx_table(doc, items, Pt, WD_TABLE_ALIGNMENT):
    h = doc.add_paragraph()
    h.paragraph_format.first_line_indent = Pt(0)
    h.add_run('引用清单').bold = True
    table = doc.add_table(rows=1, cols=5, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(['段落', '标准/条文', '审计', '版本', '类型']):
        table.rows[0].cells[idx].text = header
    for item in items[:80]:
        row = table.add_row().cells
        code = item.get('official_code') or item.get('code')
        row[0].text = f"P{item.get('para')}"
        row[1].text = f"{code} §{item.get('clause_ref')}"
        row[2].text = item.get('audit_status', '')
        row[3].text = item.get('version_status', '')
        row[4].text = item.get('clause_type', '')


def append_citation_audit_summary_to_docx(docx_path, entries, output_path=None):
    """Append a compact citation audit summary to a copy of a docx."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    summary = citation_audit_summary(entries)
    output_path = _citation_audit_docx_output_path(docx_path, output_path)
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    doc = _Document(docx_path)
    _add_citation_audit_docx_header(doc, summary, Pt, WD_ALIGN_PARAGRAPH)

    review_items = [item for item in summary['items'] if item.get('review_reasons')]
    _add_review_items_docx_table(doc, review_items, Pt, WD_TABLE_ALIGNMENT)

    _add_citation_items_docx_table(doc, summary['items'], Pt, WD_TABLE_ALIGNMENT)

    doc.save(output_path)
    return {'docx': output_path, 'summary': summary}


# ============================================================
#  输出层 2: resolve_for_chapter — 新建管线精确查找
# ============================================================

def resolve_for_chapter(chapter, topic, kb=None):
    """给定章节和专题,返回该专题应引用的规范与条款候选。"""
    if kb is None:
        kb = KB()
    from retrieval_core import RetrievalCore

    ch_info = CHAPTER_MAP.get(chapter, {})
    codes = ch_info.get('topic_standards', {}).get(topic, [])
    response = RetrievalCore(kb).match({
        'mode': 'chapter_recommend',
        'constraints': {'chapter': chapter, 'topic': topic, 'codes': codes},
        'limits': {'max_clauses': 3},
    })
    return [
        {
            'code': item['code'],
            'name': item['name'],
            'clauses': item.get('clauses', []),
            'citations': item.get('citations', []),
        }
        for item in response['results']
    ]


def _support_guarded_suggestion_candidates(results):
    usable, review, _blocked = partition_by_support_action(results)
    candidates = []
    for item in usable + review:
        raw = dict(item.get('raw') or item)
        action = item.get('support_action') or raw.get('support_action', '')
        judgment = item.get('support_judgment') or raw.get('support_judgment', '')
        signals = item.get('support_signals') or raw.get('support_signals', {})
        raw.setdefault('support_action', action)
        raw.setdefault('support_judgment', judgment)
        raw.setdefault('support_signals', signals)
        if action and action != 'use_as_evidence':
            raw['suggestion_review'] = action
        candidates.append(raw)
    return candidates


def suggest_citations(docx_path, kb=None, max_paras=50):
    """发现模式:扫描无引用编号的技术段落,搜索 KB 建议应引规范。"""
    if kb is None:
        kb = KB()
    from retrieval_core import RetrievalCore

    doc = _Document(docx_path)
    retrieval = RetrievalCore(kb)
    suggestions = []
    count = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        response = retrieval.match({
            'mode': 'citation_discovery',
            'query': text,
            'limits': {'max_results': 5, 'max_words': 6},
        })
        candidates = _support_guarded_suggestion_candidates(response['results'])
        if candidates:
            suggestions.append((i, text[:200], candidates))
            count += 1
            if count >= max_paras:
                break
    return suggestions

if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    parser = argparse.ArgumentParser(description='KB Auditor — 规范引用审计引擎')
    parser.add_argument('docx', nargs='?', help='Path to docx (audit mode)')
    parser.add_argument('--chapter', help='Chapter for resolve mode (e.g. ch05)')
    parser.add_argument('--topic', help='Topic for resolve mode (e.g. paving)')
    parser.add_argument('--self-test', action='store_true', help='Run built-in self-tests')
    parser.add_argument('--suggest', action='store_true', help='Discover mode: find uncited paragraphs, suggest KB standards')
    parser.add_argument('--all', action='store_true', help='Unified audit: verify + discover in one report')
    parser.add_argument('--summary-out', help='Output base path for citation audit summary (.json and .md)')
    parser.add_argument('--append-summary-docx', nargs='?', const='', help='Append citation audit summary to a docx copy; optional output path')
    parser.add_argument('--resolve', nargs=4, metavar=('CODE','CLAUSE','CLAIMED','HINT'),
                       help='Direct engine call: code clause claimed hint')
    args = parser.parse_args()

    if args.self_test:
        kb = KB()
        errors = [0]  # mutable counter

        def check(name, cond, detail=''):
            if not cond:
                errors[0] += 1
            print(f'  [{"PASS" if cond else "FAIL"}] {name}{" — "+detail if detail else ""}')

        # ============================================
        #  L1: 引擎正确性 — 8种值类型 × 4种代码格式
        # ============================================
        print('=' * 60)
        print('L1: 引擎正确性 (值匹配 × 代码格式 × 输出一致性)')
        print('=' * 60)

        # 1a. 值类型：带单位精确匹配
        cases_l1 = [
            # (code, clause, claimed, context, should_match, label)
            ('GB50209', '6.3.10', '不大于10mm', '踏步 宽度', True, 'mm单位'),
            ('GB50203', '5.3.3', '大于10mm', '轴线位移', True, '表格允许偏差'),
            ('CJJ82', '4.6.1', '不应低于95%', '成活率', True, '%成活率'),
            ('JGJ79', '4.3.2', '大于200mm', '分层铺填', True, '分层厚度'),
        ]
        for code, clause, claimed, ctx, expected, label in cases_l1:
            res = _resolve_citation(kb, code, clause, claimed, {'context': ctx, 'search_terms': f'{claimed} {ctx}'})
            check(f'{label}: {code} {clause} = {claimed}', res['in_clause'] == expected,
                  f'got in_clause={res["in_clause"]} detail={res["detail"]}')

        # 1b. 值不在条款中
        res = _resolve_citation(kb, 'CJJ/T 287-2018', '5.2.10', '低于5℃', '防寒 冬季')
        check('条款号错配: CJJ/T287 5.2.10≠防寒',
              not res['in_clause'],
              f'suggestion={res.get("suggestion","")[:60]}')

        # 1c. 代码格式变体
        fmt_cases = [
            ('GB 50209-2010', '6.1.5', '不应少于7d', '养护'),
            ('GB/T 10801.1-2025', '表1', '≥100kPa', '压缩'),
            ('JGJ 79-2012', '4.2.1', '不大于5%', '有机质'),
            ('CJJ/T 287-2018', '5.2.5', '不少于4次', '绿篱修剪'),
        ]
        for code, clause, claimed, ctx in fmt_cases:
            res = _resolve_citation(kb, code, clause, claimed, {'context': ctx, 'search_terms': f'{claimed} {ctx}'})
            check(f'代码格式: {code[:20]}', res.get('clause_exists') is not None,
                  f'in_clause={res["in_clause"]}')

        # 1d. 适配过滤
        res = _resolve_citation(kb, 'JGJ79', '4.3.2', '大于0.5m',
                                '分层铺填200~300mm，本工程保守取500mm')
        check('适配过滤: 本工程保守取',
              res['in_clause'] and 'adapted' in res.get('detail', ''))

        # 1e. 输出一致性：同一输入两次调用结果相同
        res_a = _resolve_citation(kb, 'GB50209', '6.1.5', '不应少于7d', '养护')
        res_b = _resolve_citation(kb, 'GB50209', '6.1.5', '不应少于7d', '养护')
        check('幂等性: 两次调用一致',
              res_a['in_clause'] == res_b['in_clause'])

        # ============================================
        #  L2: 边界条件 — 9种边界场景
        # ============================================
        print('\n' + '=' * 60)
        print('L2: 边界条件 (章节感知 × 噪声 × 分类 × 双出口)')
        print('=' * 60)

        # 2a. 章节归属
        chapter_tests = [
            ('GB50209', 'ch05', True, '铺装→ch05'),
            ('GB50203', 'ch05', True, '砌体→ch05'),
            ('CJJT287', 'ch07', False, '养护标准→不应在ch07安全章'),
            ('GB50300', 'ch06', True, '验收统一标准→ch06质量章'),
            ('JGJ59', 'ch07', True, '安全检查→ch07'),
            ('GB12523', 'ch08', True, '噪声→ch08'),
            ('GB50204', 'ch03', False, '混凝土验收→ch03不在预期(应在ch05/ch06)'),
        ]
        for code, ch, expected, label in chapter_tests:
            ok, exp_ch, reason = _check_chapter_fit(code, ch)
            check(f'章节: {label}', ok == expected,
                  f'expected={expected} got={ok} exp_ch={exp_ch}')

        # 2b. 值匹配精度验证
        r = _resolve_citation(kb, 'CJJ/T 287-2018', '5.2.10', '低于5℃')
        check('精度: 5℃不在虫害条款中', not r['in_clause'],
              f'in_clause={r["in_clause"]}')
        r = _resolve_citation(kb, 'GB50209', '6.1.5', '不应少于7d')
        check('精度: 7d在养护条款中', r['in_clause'],
              f'in_clause={r["in_clause"]}')

        # 2c. 噪声：不存在的规范
        res = _resolve_citation(kb, 'GB99999', '1.0.1', '≤100mm', '测试')
        check('不存在规范: GB99999', not res['in_clause'] and res['clause_exists'] is False,
              f'clause_exists={res["clause_exists"]}')

        # 2d. 噪声：无条款号的引用
        res = _resolve_citation(kb, 'GB50209', '不存在的条款', '≥7d', '测试')
        check('不存在条款', not res['in_clause'] and res['clause_exists'] is False,
              f'clause_exists={res["clause_exists"]}')

        # 2e. resolve_for_chapter 双出口
        results = resolve_for_chapter('ch05', 'paving', kb)
        check('resolve_for_chapter ch05/paving 返回结果', len(results) > 0,
              f'{len(results)} standards')
        if results:
            has_clauses = any(r.get('clauses') for r in results)
            has_name = any(r.get('name') for r in results)
            check('resolve_for_chapter ch05/paving 返回规范', has_name,
                  f'{len(results)} standards, clauses={has_clauses}')

        # 2f. adapt过滤不误杀普通引用
        res = _resolve_citation(kb, 'GB50209', '6.1.5', '不应少于7d',
                                '铺设后表面应覆盖湿润养护时间不应少于7d')
        check('adapt过滤不误杀: 普通"7d"养护', res['in_clause'],
              f'in_clause={res["in_clause"]}')

        # ============================================
        #  L3: 深度 — 全量审计 + 管线一致性
        # ============================================
        print('\n' + '=' * 60)
        print('L3: 深度压测 (真实方案全量 × 管线一致性 × 双出口)')
        print('=' * 60)

        # 真实方案样本：优先取环境变量 KB_AUDIT_SAMPLE_DOCX 指向的文档，
        # 否则回退到桌面上最新的、未经修复(.bak/fixed)的 .docx。
        import glob as _g
        docx = os.environ.get('KB_AUDIT_SAMPLE_DOCX', '').strip()
        if not docx:
            candidates = [
                f for f in _g.glob(os.path.join(os.path.expanduser('~'), 'Desktop', '*.docx'))
                if not f.endswith('.bak') and 'fixed' not in os.path.basename(f)
            ]
            candidates.sort(key=os.path.getmtime, reverse=True)
            docx = candidates[0] if candidates else ''
        if docx and os.path.exists(docx):
            lines, entries = audit_report(docx, kb)
            total = len(entries)
            ok_count = sum(1 for e in entries if e['resolution']['in_clause'])
            check(f'全量审计: {total}处引用', total > 0)
            if total > 0:
                match_rate = ok_count / total
                check(f'条款匹配率 ≥ 60%', match_rate >= 0.6,
                      f'{ok_count}/{total} = {100*match_rate:.0f}%')

            # 章节分布
            ch_dist = {}
            for e in entries:
                ch_dist[e['chapter']] = ch_dist.get(e['chapter'], 0) + 1
            check('审计覆盖≥3个章节', len(ch_dist) >= 3,
                  f'{len(ch_dist)} chapters: {dict(ch_dist)}')

            # 零 high severity（audit 不产生门禁级别）
            has_high = any('fabrication' in e['resolution'].get('suggestion', '')
                          for e in entries)
            check('审计零 high severity', not has_high,
                  'found possible fabrication' if has_high else '')

            # 输出审计表
            for line in lines:
                print(line)
        else:
            print('[SKIP] 真实方案样本未找到，跳过 L3 全量文档审计 (可设 KB_AUDIT_SAMPLE_DOCX)')

        # 3b. 管线一致性：audit_report 和 _resolve_citation 同一输入
        if os.path.exists(docx):
            # 取审计条目和单独解析对比（带完整上下文）
            for e in entries[:3]:
                res2 = _resolve_citation(kb, e['code'], e['clause_ref'],
                                        e['claimed'])
                consistent = e['resolution']['in_clause'] == res2['in_clause']
                check(f'审计/引擎一致性 P{e["para"]}', consistent,
                      f'audit={e["resolution"]["in_clause"]} engine={res2["in_clause"]}')

        # 3c. resolve_for_chapter 覆盖全部章节
        all_chapters = [k for k in CHAPTER_MAP.keys()
                       if CHAPTER_MAP[k].get('topic_standards')]
        covered = 0
        for ch in all_chapters:
            for topic in CHAPTER_MAP[ch]['topic_standards']:
                r = resolve_for_chapter(ch, topic, kb)
                if r and any(x.get('name') for x in r):
                    covered += 1
                    break
        check(f'章节覆盖: {covered}/{len(all_chapters)} 章有规范映射',
              covered >= len(all_chapters) * 0.8)

        # ════════════════════════════════════════════
        #  规则5合规检查：全部验证函数代码纯净度
        guarded_candidates = _support_guarded_suggestion_candidates([
            {'raw': {'file': 'blocked'}, 'support_action': 'block_forbidden'},
            {'raw': {'file': 'warned'}, 'support_action': 'warn_insufficient_support', 'support_judgment': 'insufficient_support'},
            {'raw': {'file': 'evidence'}, 'support_action': 'use_as_evidence', 'support_judgment': 'supported'},
        ])
        check('support_guard suggestion gate blocks forbidden', all(c.get('file') != 'blocked' for c in guarded_candidates))
        check('support_guard suggestion gate keeps review marker', any(c.get('suggestion_review') == 'warn_insufficient_support' for c in guarded_candidates))
        check('support_guard suggestion gate keeps evidence', any(c.get('file') == 'evidence' for c in guarded_candidates))

        #  _build_ai_hint/_strip_cmp 等辅助函数白名单（允许关键词/比较词处理）
        # ════════════════════════════════════════════
        import inspect as _inspect
        _white = {'_build_ai_hint', '_strip_cmp', '_match', 'check'}
        _violations = []
        _FORBIDDEN = [
            (r're\.findall\(r?[\x27\x22][\u4e00-\u9fff]', 'keyword extraction'),
            (r"(?:'quantitative'|'reference').*return", 'semantic classification'),
            (r'other_std|cross_std|_cross_std_search|_classify_context', 'cross-standard inference'),
        ]
        for _name, _fn in [(n, getattr(sys.modules[__name__], n))
                          for n in dir(sys.modules[__name__])
                          if n.startswith('_') and callable(getattr(sys.modules[__name__], n, None))]:
            if _name in _white: continue
            try:
                _src = _inspect.getsource(_fn)
            except (TypeError, OSError):
                continue
            for _pat, _desc in _FORBIDDEN:
                if re.search(_pat, _src):
                    _violations.append(f'{_desc} in {_name}()')
        for v in _violations:
            check(f'RULE5: {v}', False, f'SKILL.md rule5 violation')
            errors[0] += 1 if len(_violations) == 0 else len(_violations)

        # ============================================
        #  汇总
        # ============================================
        print('\n' + '=' * 60)
        if errors[0] == 0:
            print(f'ALL CHECKS PASSED')
        else:
            print(f'{errors[0]} FAILURES')
            sys.exit(1)

    elif args.all and args.docx:
        v_lines, entries = audit_report(args.docx)
        for line in v_lines:
            print(line)
        if args.summary_out:
            written = write_citation_audit_summary(args.docx, entries, args.summary_out)
            print(f"\nCitation audit summary: {written['md']}")
        if args.append_summary_docx is not None:
            output_path = args.append_summary_docx or None
            written_docx = append_citation_audit_summary_to_docx(args.docx, entries, output_path)
            print(f"Citation audit docx: {written_docx['docx']}")
        suggestions = suggest_citations(args.docx)
        print(f'\n=== 发现模式：{len(suggestions)} 处无引用技术段落 ===')
        for s in suggestions[:10]:
            print(f'\nP{s[0]}: {s[1][:120]}...')
            for r in s[2][:3]:
                fc = _ec(r.get('file',''))
                name, _, _ = kb.get_name(fc) if fc else (None, 0, '')
                h = r.get('heading','')
                cm = re.search(r'(\d+\.\d+(?:\.\d+)?)', h)
                cl = cm.group(1) if cm else h[:30]
                review = r.get('suggestion_review') or r.get('support_action', '')
                review_note = f' [review:{review}]' if review and review != 'use_as_evidence' else ''
                print(f'  → {fc} §{cl} ({name}){review_note}')

    elif args.suggest and args.docx:
        suggestions = suggest_citations(args.docx)
        print(f'=== 发现模式：{os.path.basename(args.docx)} ===')
        print(f'共 {len(suggestions)} 处无引用技术段落\n')
        for pi, text, results in suggestions:
            print(f'P{pi}: {text[:150]}...')
            for r in results[:3]:
                f = r.get('file', '')[:50]
                h = r.get('heading', '')[:50]
                review = r.get('suggestion_review') or r.get('support_action', '')
                review_note = f' [review:{review}]' if review and review != 'use_as_evidence' else ''
                print(f'  → {f} | {h}{review_note}')
            print()

    elif args.resolve:
        code, clause, claimed, hint = args.resolve
        res = _resolve_citation(KB(), code, clause, claimed,
                                {'search_terms': f'{claimed} {hint}', 'context': hint})
        print(json.dumps(res, ensure_ascii=False, indent=2))

    elif args.chapter and args.topic:
        results = resolve_for_chapter(args.chapter, args.topic)
        print(json.dumps(results, ensure_ascii=False, indent=2))

    elif args.docx:
        lines, entries = audit_report(args.docx)
        for line in lines:
            print(line)
        if args.summary_out:
            written = write_citation_audit_summary(args.docx, entries, args.summary_out)
            print(f"\nCitation audit summary: {written['md']}")
        if args.append_summary_docx is not None:
            output_path = args.append_summary_docx or None
            written_docx = append_citation_audit_summary_to_docx(args.docx, entries, output_path)
            print(f"Citation audit docx: {written_docx['docx']}")

    else:
        print('Usage: python kb_auditor.py <docx> | --chapter ch05 --topic paving | --self-test')
