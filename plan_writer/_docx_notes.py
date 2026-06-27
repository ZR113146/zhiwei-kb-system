#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Shared docx annotation helpers for the plan_writer enhancer/corrector.

Single source of truth for the yellow KB review note style and paragraph
lookup, so the enhancer and corrector cannot drift apart.
"""

from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor


def yellow_append(paragraph, text, size=8):
    """Append a yellow-highlighted KB review note run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)
    run.font.size = Pt(size)
    return run


def paragraph_by_index(doc, para_idx):
    """Return doc.paragraphs[para_idx] or None when out of range."""
    if para_idx < 0 or para_idx >= len(doc.paragraphs):
        return None
    return doc.paragraphs[para_idx]
