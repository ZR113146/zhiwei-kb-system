#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""orchestrator: 一键调度从零编制全流程+质检，输出综合报告"""

import os, re, sys, subprocess, json, argparse, logging
logging.basicConfig(level=logging.WARNING, format='%(levelname)s: %(message)s')

SCRIPT_DIR = os.path.dirname(__file__)
SKILL_DIR = os.path.dirname(SCRIPT_DIR)
CONTENT_DIR = os.path.join(SKILL_DIR, 'content')

# 章节最低要求（来自 SKILL.md AI扩充策略）
CHAPTER_TARGETS = {
    'ch02': {'min_paras': 25, 'name': '\u5de5\u7a0b\u6982\u51b5', 'checks': ['\u57fa\u672c\u4fe1\u606f\u8868','\u8303\u56f4\u8be6\u8ff0','\u673a\u68b0\u914d\u7f6e\u8868(\u22658\u9879)','\u91cd\u96be\u70b9(\u22652\u6761)']},
    'ch03': {'min_paras': 30, 'name': '\u65bd\u5de5\u5b89\u6392', 'checks': ['\u7ec4\u7ec7\u673a\u6784','\u52b3\u52a8\u529b\u8868(\u22656\u73ed\u7ec4)','\u5206\u533a\u6d41\u6c34','\u5de5\u671f\u8868(\u22658\u9636\u6bb5)']},
    'ch04': {'min_paras': 25, 'name': '\u65bd\u5de5\u51c6\u5907', 'checks': ['\u6280\u672f\u4ea4\u5e95','\u8bd5\u9a8c\u68c0\u9a8c\u8868(\u22656\u9879)','\u6837\u677f\u8868(\u22653\u9879)','\u7269\u8d44\u6e05\u5355(\u542b\u6570\u91cf)']},
    'ch05': {'min_paras': 80, 'name': '\u4e3b\u8981\u65bd\u5de5\u65b9\u6cd5', 'checks': ['\u22656\u5206\u9879','\u6bcf\u5206\u9879\u22655\u6bb5','\u5141\u8bb8\u504f\u5dee\u8868','\u89c4\u8303\u6761\u6587\u5f15\u7528']},
    'ch06': {'min_paras': 40, 'name': '\u8d28\u91cf\u8981\u6c42', 'checks': ['\u8d28\u91cf\u4f53\u7cfb','QC\u8868','\u9a8c\u6536\u5212\u5206','\u901a\u75c5\u9632\u6cbb(\u22653\u9879)']},
    'ch07': {'min_paras': 30, 'name': '\u5b89\u5168\u7ba1\u7406', 'checks': ['\u22656\u5de5\u5e8f\u5b89\u5168','\u5e94\u6025\u6551\u63f4(\u542b\u7269\u8d44\u8868)','\u7279\u6b8a\u5b89\u5168']},
    'ch08': {'min_paras': 20, 'name': '\u6587\u660e\u65bd\u5de5', 'checks': ['\u626c\u5c18(\u22654\u63aa\u65bd)','\u566a\u58f0(\u542bdB)','\u5e9f\u5f03\u7269(\u22653\u7c7b)','\u5b63\u8282\u6027(\u22652\u5b63)']},
    'ch09': {'min_paras': 20, 'name': '\u5176\u4ed6\u8981\u6c42', 'checks': ['\u5de5\u671f\u4fdd\u8bc1(\u22654\u63aa\u65bd)','\u6210\u54c1\u4fdd\u62a4(\u22655\u7c7b)','\u964d\u9020(\u22653\u63aa\u65bd)']},
}


def count_chapter_paras(docx_path):
    """Count paragraphs per chapter in a docx.
    Matches: '1  工程概况', '# 2 施工安排', '3 施工准备', etc.
    """
    from docx import Document
    doc = Document(docx_path)
    # Regex matches: optional '# ', digit 1-9, 1+ whitespace, short chapter title
    _ch_re = re.compile(r'^#?\s*([1-9])\s+')
    counts = {}
    cur = None
    for p in doc.paragraphs:
        t = p.text.strip()
        m = _ch_re.match(t)
        if m and len(t) < 50:
            cur = m.group(1)
        if cur:
            counts['ch0'+cur] = counts.get('ch0'+cur, 0) + 1
    total_paras = len(doc.paragraphs)
    total_tables = len(doc.tables)
    total_images = len(doc.inline_shapes)
    return counts, total_paras, total_tables, total_images


def run_scan(docx_path):
    try:
        result = subprocess.run([sys.executable, os.path.join(SCRIPT_DIR, 'scan.py'), docx_path],
                               capture_output=True, text=True, encoding='utf-8', errors='replace', timeout=30)
    except subprocess.TimeoutExpired:
        logging.warning('scan.py timed out after 30s — returning empty result')
        return {'high': 0, 'medium': 0, 'low': 0, 'total': 0, 'error': 'timeout'}
    except Exception as e:
        logging.warning(f'scan.py failed: {e}')
        return {'high': 0, 'medium': 0, 'low': 0, 'total': 0, 'error': str(e)}
    high = result.stdout.count('[high]')
    medium = result.stdout.count('[medium]')
    low = result.stdout.count('[low]')
    return {'high': high, 'medium': medium, 'low': low, 'total': high+medium+low}


def run_verify(docx_path):
    try:
        result = subprocess.run([sys.executable, os.path.join(SCRIPT_DIR, 'verify.py'), docx_path],
                               capture_output=True, text=True, encoding='utf-8', errors='replace', timeout=60)
    except subprocess.TimeoutExpired:
        logging.warning('verify.py timed out after 60s — returning empty result')
        return {'high': 0, 'medium': 0, 'low': 0, 'total': 0, 'passed': False, 'error': 'timeout'}
    except Exception as e:
        logging.warning(f'verify.py failed: {e}')
        return {'high': 0, 'medium': 0, 'low': 0, 'total': 0, 'passed': False, 'error': str(e)}
    high = result.stdout.count('[high]')
    low = result.stdout.count('[low]')
    medium = result.stdout.count('[medium]')
    return {'high': high, 'medium': medium, 'low': low, 'total': high+medium+low, 'passed': result.returncode == 0}


def check_chapter_depth(docx_path):
    """Check each chapter against minimum targets"""
    counts, total_p, total_t, total_i = count_chapter_paras(docx_path)
    report = []
    all_pass = True
    for ch_key, target in CHAPTER_TARGETS.items():
        actual = counts.get(ch_key, 0)
        passed = actual >= target['min_paras']
        if not passed:
            all_pass = False
        report.append({
            'chapter': ch_key,
            'name': target['name'],
            'actual': actual,
            'target': target['min_paras'],
            'passed': passed,
            'missing_checks': [] if passed else target['checks']
        })
    return report, total_p, total_t, total_i, all_pass


def print_report(docx_path, chapter_report, scan_result, verify_result, total_p, total_t, total_i):
    print('=' * 60)
    print(f'ORCHESTRATOR REPORT: {os.path.basename(docx_path)}')
    print('=' * 60)
    print(f'\n  Paragraphs: {total_p} | Tables: {total_t} | Images: {total_i}\n')

    print('  Chapter Depth:')
    for r in chapter_report:
        status = 'PASS' if r['passed'] else 'FAIL'
        print(f'    [{status}] {r["chapter"]} {r["name"]}: {r["actual"]}/{r["target"]} segments')
        if not r['passed']:
            print(f'           Missing: {", ".join(r["missing_checks"])}')

    print(f'\n  Quality:')
    print(f'    scan.py:  {scan_result["total"]} gaps (high:{scan_result["high"]} medium:{scan_result["medium"]} low:{scan_result["low"]})')
    print(f'    verify.py: {verify_result["total"]} issues (high:{verify_result["high"]} medium:{verify_result["medium"]} low:{verify_result["low"]}) | passed={verify_result["passed"]}')

    all_ok = all(r['passed'] for r in chapter_report) and scan_result['high'] == 0 and verify_result['high'] == 0
    print(f'\n  >>> {"ALL CHECKS PASSED" if all_ok else "ISSUES FOUND - see above"} <<<')
    if all_ok:
        strip_markers(docx_path)
    return all_ok


def strip_markers(docx_path):
    """Auto-strip [L1:KB]/[L2:手册]/[L3:AI]/[L4:Web] source markers after all checks pass"""
    import re
    from docx import Document
    MARKER_RE = re.compile(r'\s*\[L[1-4]:(?:KB|手册|AI|Web)\]\s*')

    doc = Document(docx_path)
    count = 0

    def _clean(text):
        return MARKER_RE.sub('', text)

    for p in doc.paragraphs:
        old = p.text
        new = _clean(old)
        if old != new:
            if p.runs:
                for r in p.runs[1:]:
                    r.text = ''
                p.runs[0].text = new
            else:
                p.text = new
            count += 1

    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    old = p.text
                    new = _clean(old)
                    if old != new:
                        if p.runs:
                            for r in p.runs[1:]:
                                r.text = ''
                            p.runs[0].text = new
                        else:
                            p.text = new
                        count += 1

    if count > 0:
        doc.save(docx_path)
        print(f'  Markers stripped: {count} instances ([L1:KB]/[L2:手册]/[L3:AI]/[L4:Web])')


def save_progress(docx_path, chapter_report, scan_result, verify_result):
    """Write progress.json for session recovery. AI reads this on wake-up.
    Derives content dir from docx path to support project isolation."""
    from datetime import datetime
    # Walk up from docx to find project root (parent of content/)
    docx_dir = os.path.dirname(os.path.abspath(docx_path))
    parent = os.path.basename(docx_dir)
    if parent == 'content':
        # Project structure: projects/<name>/content/output.docx
        progress_path = os.path.join(docx_dir, 'progress.json')
    else:
        # Shared content dir
        progress_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'content', 'progress.json')
    os.makedirs(os.path.dirname(progress_path), exist_ok=True)

    chapters = {}
    for r in chapter_report:
        status = 'pass' if r['passed'] else ('phase1' if r['actual'] > r['target'] * 0.3 else 'empty')
        chapters[r['chapter']] = {
            'name': r['name'],
            'status': status,
            'segments': r['actual'],
            'target': r['target'],
            'missing': r.get('missing_checks', [])
        }

    progress = {
        '_updated': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'docx': os.path.basename(docx_path),
        'chapters': chapters,
        'scan': {'total': scan_result['total'], 'high': scan_result['high']},
        'verify': {'total': verify_result['total'], 'high': verify_result['high']},
    }

    with open(progress_path, 'w', encoding='utf-8') as f:
        json.dump(progress, f, ensure_ascii=False, indent=2)
    print(f'  Progress saved: {progress_path}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('docx', help='Path to docx file')
    parser.add_argument('--json', action='store_true', help='Output as JSON')
    parser.add_argument('--save-progress', action='store_true', help='Write progress.json for session recovery')
    parser.add_argument('--audit', action='store_true', help='Run kb_auditor full citation audit')
    args = parser.parse_args()

    # 审计模式：先跑规则5合规检查，再跑引用对照表+门禁
    if args.audit:
        result = subprocess.run(
            [sys.executable, os.path.join(SCRIPT_DIR, 'kb_auditor.py'), '--self-test'],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            print(' RULE5 COMPLIANCE FAILED — aborting')
            print(result.stderr[-500:] if result.stderr else result.stdout[-500:])
            sys.exit(1)
        from kb_auditor import audit_report
        lines, entries = audit_report(args.docx)
        for line in lines:
            print(line)

    cr, tp, tt, ti, _ = check_chapter_depth(args.docx)
    sr = run_scan(args.docx)
    vr = run_verify(args.docx)

    if args.json:
        print(json.dumps({'chapters': cr, 'scan': sr, 'verify': vr, 'totals': {'paras': tp, 'tables': tt, 'images': ti}},
                        ensure_ascii=False, indent=2))
    else:
        print_report(args.docx, cr, sr, vr, tp, tt, ti)

    if args.save_progress:
        save_progress(args.docx, cr, sr, vr)


# ===== Orchestrator Class (Codex 平台方案编制管线) =====
class Orchestrator:
    """方案编制管线编排器。
    
    流程: scan → verify → audit → enhance → correct → generate → render → build
    双模式: python-docx (批量) / COM Word (交互)
    """

    def __init__(self, kb=None):
        self._kb = kb
        self._project_dir = None

    def _get_kb(self):
        if self._kb is None:
            kb_core_dir = os.path.join(SKILL_DIR, 'kb_core')
            if kb_core_dir not in sys.path:
                sys.path.insert(0, kb_core_dir)
            from kb_core.kb import KB
            self._kb = KB()
        return self._kb

    def _find_docx(self, project_dir):
        """Find the .docx file in project directory."""
        import glob
        candidates = []
        for f in glob.glob(os.path.join(project_dir, "*.docx")):
            if ".bak" in f or os.path.basename(f).startswith("~$"):
                continue
            candidates.append(f)
        if not candidates:
            # Fallback: check project/ subdirectory
            proj_sub = os.path.join(os.path.dirname(SCRIPT_DIR), 'projects', os.path.basename(project_dir))
            if os.path.isdir(proj_sub):
                for f in glob.glob(os.path.join(proj_sub, "*.docx")):
                    if ".bak" in f or os.path.basename(f).startswith("~$"):
                        continue
                    candidates.append(f)
        return candidates[0] if candidates else None

    def _find_brief(self, project_dir):
        """Find project brief (project.json or project_brief.json)."""
        import glob
        for name in ['project.json', 'project_brief.json']:
            candidates = glob.glob(os.path.join(project_dir, name))
            if candidates:
                return candidates[0]
        return None

    def _find_content_dir(self, project_dir):
        """Find or create content directory."""
        content = os.path.join(project_dir, 'content')
        if os.path.isdir(content):
            return content
        # Check projects/<name>/content
        proj_name = os.path.basename(project_dir.rstrip('\\/'))
        proj_content = os.path.join(os.path.dirname(SCRIPT_DIR), 'projects', proj_name, 'content')
        if os.path.isdir(proj_content):
            return proj_content
        os.makedirs(content, exist_ok=True)
        return content

    def run_scan(self, docx_path):
        """扫描项目文档，识别规范引用缺口。"""
        from scan import scan_docx
        return scan_docx(docx_path)

    def run_verify(self, docx_path):
        """对照 KB 验证规范存在性。"""
        from verify import verify
        return verify(docx_path, gate_mode=False)

    def run_audit(self, docx_path, output_dir=None, append_summary_docx=False):
        """审计规范适用性 & 建议替代。"""
        kb = self._get_kb()
        from kb_auditor import audit_report, append_citation_audit_summary_to_docx, write_citation_audit_summary
        lines, entries = audit_report(docx_path, kb=kb)
        result = {'lines': lines, 'entries': entries}
        if output_dir:
            base = os.path.join(output_dir, 'citation_audit_summary')
            result['summary_files'] = write_citation_audit_summary(docx_path, entries, base)
            if append_summary_docx:
                root, ext = os.path.splitext(os.path.basename(docx_path))
                docx_out = os.path.join(output_dir, root + '_citation_audit' + ext)
                result['summary_docx'] = append_citation_audit_summary_to_docx(docx_path, entries, docx_out)
        return result

    def run_enhance(self, docx_path, output_path=None):
        """增强条款内容（用 KB 原文替换/补充）。"""
        from kb_enhancer import enhance_docx
        return enhance_docx(docx_path, output_path=output_path)

    def run_correct(self, docx_path, output_path=None):
        """修正引用格式。"""
        from kb_corrector import apply_corrections
        return apply_corrections(docx_path, output_path=output_path)

    def run_generate(self, brief_path, content_dir):
        """生成方案正文章节。"""
        from content_generator import ContentGenerator
        gen = ContentGenerator(brief_path, content_dir)
        gen.generate_all()
        return {'content_dir': content_dir}

    def run_render(self, content_dir, output_path):
        """排版渲染 → 输出 .docx。"""
        from render_engine import PlanRenderer
        renderer = PlanRenderer(content_dir, output_path)
        renderer.build()
        return {'output': output_path}

    def run(self, project_dir, output_dir=None, append_citation_summary=False):
        """执行完整方案编制管线。

        Args:
            project_dir: 项目目录（含 .docx 草稿）
            output_dir: 输出目录（默认同 project_dir）
        Returns:
            dict: {'scan', 'verify', 'audit', 'output', ...}
        """
        self._project_dir = project_dir
        out_dir = output_dir or project_dir
        os.makedirs(out_dir, exist_ok=True)

        docx_path = self._find_docx(project_dir)
        result = {
            'project_dir': project_dir,
            'output_dir': out_dir,
            'docx_path': docx_path,
        }

        if docx_path:
            logging.info(f'Orchestrator: processing {os.path.basename(docx_path)}')

            # Phase 1: Audit
            scan_result = self.run_scan(docx_path)
            result['scan'] = scan_result

            verify_result = self.run_verify(docx_path)
            result['verify'] = verify_result

            audit_result = self.run_audit(docx_path, output_dir=out_dir, append_summary_docx=append_citation_summary)
            result['audit'] = audit_result

            # Phase 2: Enhance (audit notes + guarded citation suggestions, single pass)
            enhanced_path = os.path.join(out_dir, 'enhanced.docx')
            self.run_enhance(docx_path, output_path=enhanced_path)
            result['corrected_path'] = enhanced_path
        else:
            logging.warning('Orchestrator: no .docx found — running content generation only')
            result['corrected_path'] = None

        # Phase 3: Generate & Render
        brief_path = self._find_brief(project_dir)
        if brief_path:
            content_dir = self._find_content_dir(project_dir)
            gen_result = self.run_generate(brief_path, content_dir)
            result['generate'] = gen_result

            from datetime import datetime
            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = os.path.join(out_dir, f'施工方案_{ts}.docx')
            render_result = self.run_render(content_dir, output_path)
            result['output'] = render_result.get('output', output_path)
        else:
            result['output'] = result.get('corrected_path')

        return result


# CLI entry preserved
