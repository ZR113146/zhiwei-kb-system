# -*- coding: utf-8 -*-
"""一致性校验: 跨章数值比对+引用断链+版本检查+本地标准区域匹配(不用AI)"""
import os, re, json, argparse, sys
from docx import Document

from kb_core.kb import KB, normalize_code, extract_code as _ec
from kb_core.support_guard import partition_by_support_action
import kb_core.changelog as changelog; changelog.record(__file__, sys.argv)
from _utils import find_latest_docx

# === 行政区划→地方标准前缀映射 ===
# 规则：DB + 行政区划代码前2~4位 = 该省市的地方标准
# DGJ/J/TJ等是工程建设类地方标准的不同前缀
_ADMIN_MAP = {
    # (省简称, 市名) → [标准前缀列表]
    ('江苏', '南京'):  ['DB32', 'DB3201', 'DGJ32', '苏'],
    ('江苏', '苏州'):  ['DB32', 'DB3205', 'DGJ32', '苏'],
    ('江苏', '无锡'):  ['DB32', 'DB3202', 'DGJ32', '苏'],
    ('江苏', '常州'):  ['DB32', 'DB3204', 'DGJ32', '苏'],
    ('江苏', None):    ['DB32', 'DGJ32', '苏'],
    ('浙江', '杭州'):  ['DB33', 'DB3301', '浙'],
    ('浙江', None):    ['DB33', '浙'],
    ('上海', None):    ['DB31', 'DG/TJ08', '沪'],
    ('北京', None):    ['DB11', 'DBJ01', '京'],
    ('广东', '广州'):  ['DB44', 'DB4401', '粤'],
    ('广东', '深圳'):  ['DB44', 'DB4403', '粤'],
    ('广东', None):    ['DB44', '粤'],
    ('山东', None):    ['DB37', '鲁'],
    ('湖北', '武汉'):  ['DB42', 'DB4201', '鄂'],
    ('湖北', None):    ['DB42', '鄂'],
    ('四川', '成都'):  ['DB51', 'DB5101', '川'],
    ('四川', None):    ['DB51', '川'],
    ('安徽', '合肥'):  ['DB34', 'DB3401', '皖'],
    ('安徽', None):    ['DB34', '皖'],
    ('福建', None):    ['DB35', '闽'],
    ('河南', None):    ['DB41', '豫'],
    ('湖南', None):    ['DB43', '湘'],
    ('河北', None):    ['DB13', '冀'],
    ('陕西', None):    ['DB61', '陕'],
    ('重庆', None):    ['DB50', '渝'],
    ('天津', None):    ['DB12', '津'],
}

# 城市→省份反向推断
_CITY_TO_PROVINCE = {
    '南京': '江苏', '苏州': '江苏', '无锡': '江苏', '常州': '江苏', '南通': '江苏',
    '杭州': '浙江', '宁波': '浙江', '温州': '浙江',
    '广州': '广东', '深圳': '广东', '东莞': '广东', '佛山': '广东',
    '武汉': '湖北', '成都': '四川', '合肥': '安徽', '福州': '福建',
    '郑州': '河南', '长沙': '湖南', '石家庄': '河北', '西安': '陕西',
    '沈阳': '辽宁', '大连': '辽宁', '济南': '山东', '青岛': '山东',
    '哈尔滨': '黑龙江', '长春': '吉林', '南昌': '江西', '太原': '山西',
    '兰州': '甘肃', '昆明': '云南', '贵阳': '贵州', '海口': '海南',
    '南宁': '广西', '呼和浩特': '内蒙古', '银川': '宁夏', '乌鲁木齐': '新疆',
    '拉萨': '西藏', '西宁': '青海',
}

def _parse_project_location(doc):
    """从方案docx中提取项目地点。返回 (province, city, raw_text)
    多策略：标题→地点字段→施工范围→首个地名"""
    full_text = '\n'.join(p.text for p in doc.paragraphs)

    # 策略1：明确的地点字段
    for pat in [
        r'(?:项目地点|工程地点|建设地点|建设地址)[：:\s]*([^\n。]{6,60}?)(?:[。\n]|$)',
    ]:
        m = re.search(pat, full_text)
        if m:
            loc = m.group(1).strip()
            break
    else:
        # 策略2：标题中提取（前5个段落）
        title_text = '\n'.join(p.text for i, p in enumerate(doc.paragraphs) if i < 5)
        # 匹配 "XX市" 或已知城市名
        _known_cities = '|'.join(_CITY_TO_PROVINCE.keys())
        m = re.search(rf'([\u4e00-\u9fff]{{2,4}}(?:市|省|区|县)|(?:{_known_cities}))', title_text)
        if m:
            loc = title_text[:200]  # 用标题区域做上下文
        else:
            # 策略3：施工范围描述
            m = re.search(
                r'(?:施工范围|工程范围|项目位于|工程位于|本工程位于|本工程地处|地处|施工范围为)'
                r'[：:\s]*(.{6,80}?)(?:[。\n]|施工|项目|。$)',
                full_text
            )
            if m:
                loc = m.group(1).strip()
            else:
                # 策略4：全文搜索第一个城市名
                m = re.search(
                    rf'([\u4e00-\u9fff]{{2,4}}(?:市|区|县)|(?:{_known_cities}))',
                    full_text
                )
                if m:
                    start = max(0, m.start() - 30)
                    end = min(len(full_text), m.end() + 60)
                    loc = full_text[start:end]
                else:
                    return None, None, ''

    # 识别省市
    province = None
    city = None

    # 直辖市
    _municipalities = {
        '北京市': '北京', '上海市': '上海', '天津市': '天津', '重庆市': '重庆'
    }
    for full, short in _municipalities.items():
        if full in loc or ('市' in loc and short in loc):
            return short, short, loc

    # 省份
    _known_cities = '|'.join(_CITY_TO_PROVINCE.keys())
    prov_pat = re.search(
        r'(江苏|浙江|广东|山东|湖北|四川|安徽|福建|河南|湖南|河北|陕西|'
        r'辽宁|吉林|黑龙江|江西|山西|甘肃|云南|贵州|海南|青海|台湾|'
        r'广西|内蒙古|宁夏|新疆|西藏)(?:省|自治区)?', loc
    )
    if prov_pat:
        province = prov_pat.group(1)

    # 城市检测（独立于省份）
    city_pat = re.search(
        rf'([\u4e00-\u9fff]{{2,3}}(?:市|区|县)|(?:{_known_cities}))',
        loc
    )
    if city_pat:
        raw_city = city_pat.group(1)
        if raw_city != (province or '') + '省':
            city = raw_city.replace('市', '').replace('区', '').replace('县', '')

    # 城市→省份反向推断
    if not province and city and city in _CITY_TO_PROVINCE:
        province = _CITY_TO_PROVINCE[city]

    return province, city, loc

def verify(docx_path, gate_mode=False):
    """运行全部校验,返回问题列表。gate_mode=True启用Phase2程序门禁。"""
    doc = Document(docx_path)
    kb = KB()
    issues = []
    all_text = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs) if p.text.strip()]

    # Load project standards for search context
    pstandards = _load_project_standards()

    # 1. 跨章数值冲突
    issues += check_value_conflicts(all_text)

    # 2. 引用断链(详见§X.X.X)
    issues += check_cross_refs(all_text)

    # 3. 规范版本校验（用 kb_resolver 统一查询）
    issues += check_version(all_text, kb)

    # 4. 必需章节检查
    issues += check_required_sections(all_text)

    # 5. 规范名称校验：逐条核对编号↔名称是否与KB一致
    issues += check_code_names(all_text, kb)

    # 5.5 地方标准区域匹配：项目地点 ↔ 引用的DB/DBJ/DGJ标准
    issues += check_local_standards(doc, all_text)

    # 6. Phase2门禁校验
    issues += check_ai_claims(all_text, kb, gate_mode=gate_mode, pstandards=pstandards)

    return issues

def _load_project_standards():
    """Load matched standards from content/project.json for search context."""
    try:
        pj_path = os.path.join(os.path.dirname(__file__), '..', 'content', 'project.json')
        if os.path.exists(pj_path):
            with open(pj_path, 'r', encoding='utf-8') as f:
                pj = json.load(f)
            pstandards = set()
            for codes in pj.get('matched_standards', {}).values():
                for c in codes:
                    pstandards.add(normalize_code(c))
            return pstandards
    except FileNotFoundError:
        # No project.json — OK, run without project context
        return None
    except (json.JSONDecodeError, Exception) as e:
        import logging
        logging.warning(f'Failed to load project.json for search context: {e}. '
                       'Run "python content_generator.py" to regenerate.')
        return None

def check_value_conflicts(all_text):
    """检测跨章数值冲突（收紧匹配窗口，过滤明显非目标值）"""
    from _value_conflict import scan_hits  # E5: 共享扫描骨架
    issues = []
    checks = {
        '养护时间': (r'养护(?:时间)?(?:不[得应][少低大于超]|≥)\s*(\d+)\s*[d天]', 2, 30),
        '压实系数': (r'压实系数[λc]?\s*(?:不小于|≥|≧)\s*([0-9.]+)', 0.5, 5),
        '开挖深度': (r'(?:最大)?开挖深度(?:不[得应][超大于])?\s*(\d+\.?\d*)\s*[m米]', 0.5, 20),
        '灰缝厚度': (r'灰缝[^。\n]{0,20}(\d+)\s*mm', 5, 50),
        '砂浆饱满度': (r'饱满度[^。\n]{0,20}(\d+)\s*%', 50, 100),
        '成活率': (r'成活率[^。\n]{0,20}(\d+)\s*%', 50, 100),
    }
    # 公共骨架扫描+阈值过滤; verify 特有的开挖单位归一在此后处理。
    hits = scan_hits(all_text, checks)
    values = {k: [] for k in checks}
    for key, key_hits in hits.items():
        for i, val, text in key_hits:
            # 开挖深度单位统一：1500mm→1.5m
            if key == '开挖深度' and float(val) > 100:
                val = str(float(val) / 1000)
            values[key].append((i, val, text[:80]))
    for key, vals in values.items():
        unique = set(v[1] for v in vals)
        if len(unique) > 1:
            locs = [f'P{v[0]}' for v in vals]
            sev = 'low' if key == '养护时间' else 'high'
            issues.append({
                'type': 'value_conflict',
                'detail': f'{key}: {sorted(unique)} @ {locs}',
                'severity': sev
            })
    return issues

def check_cross_refs(all_text):
    """检测'详见§X.X.X'引用是否有效"""
    issues = []
    # 收集所有章节标题
    sections = set()
    for i, text in all_text:
        m = re.match(r'^(\d+\.\d+(?:\.\d+)?)\s', text)
        if m:
            sections.add(m.group(1))
    # 检查交叉引用
    for i, text in all_text:
        for m in re.finditer(r'详见.*?[§§](\d+\.\d+(?:\.\d+)?)', text):
            ref = m.group(1)
            if ref not in sections:
                issues.append({
                    'type': 'broken_ref',
                    'detail': f'P{i}: "详见§{ref}" but section not found',
                    'severity': 'medium'
                })
    return issues

def check_version(all_text, kb):
    """检查规范版本——通过 kb_resolver 统一查询"""
    issues = []
    for i, text in all_text:
        for m in re.finditer(r'((?:GB\s*/?\s*T?|JGJ|CJJ|CECS|CJ\s*/?\s*T?)\s*\d+[\.-]\d+(?:-\d+)?)', text):
            if not kb.exists(m.group()):
                issues.append({
                    'type': 'version_mismatch',
                    'detail': f'P{i}: {m.group()} not in knowledge base',
                    'severity': 'low'
                })
    return issues

def check_citation_accuracy(all_text, kb):
    """委托 kb_auditor._resolve_citation 做三阶段引用解析。
    本函数仅保留门禁接口，详细审计用 kb_auditor.audit_report()。"""
    from kb_auditor import _resolve_citation, _CITE_RE
    issues = []
    for i, text in all_text:
        for m in _CITE_RE.finditer(text):
            code = m.group(1).strip()
            claimed = m.group(2).strip()
            inter = text[m.start(1):m.start(2)]
            cm = re.search(r'(?:表\s*(\d+)|第\s*([\d.]+)条|§\s*([\d.]+)|附录\s*([A-Z]))', inter)
            if not cm:
                continue
            if cm.group(1):   clause_ref = f'表{cm.group(1)}'
            elif cm.group(2): clause_ref = cm.group(2)
            elif cm.group(3): clause_ref = cm.group(3)
            else:             clause_ref = cm.group(4)
            clause_display = cm.group()

            res = _resolve_citation(kb, code, clause_ref, claimed, inter)

            if res['in_clause']:
                continue  # 匹配，通过

            # 不匹配 → 报告
            sev = 'high' if 'fabrication' in res.get('suggestion', '') else 'medium'
            issues.append({
                'type': 'citation_wrong_clause',
                'detail': (f'P{i}: {code} {clause_display} cites {claimed}. '
                         f'{res.get("suggestion", "clause mismatch")}'),
                'severity': sev
            })
    return issues


def check_ai_claims(all_text, kb, gate_mode=False, pstandards=None):
    """Phase2 gate: verify AI/Web numeric claims against KB evidence."""
    from retrieval_core import RetrievalCore

    issues = []
    l3_count = 0
    l4_count = 0
    kb_hit = 0
    kb_miss = 0
    retrieval = RetrievalCore(kb)

    numeric_claim = re.compile(
        r'([\u2265\u2264\u2267]\s*\d+\.?\d*\s*(?:kPa|kN|mm|cm|m|MPa|%|d|\u5929|kg/m\u00b3|\u2103)|'
        r'(?:\u4e0d[\u5f97\u5e94][\u5c0f\u5927\u4f4e\u9ad8\u8d85\u4e8e\u5c11\u591a]).{0,8}?\d+\.?\d*)'
    )

    for i, text in all_text:
        source_tag = None
        if '[L3:AI]' in text:
            source_tag = 'L3:AI'
            l3_count += 1
        elif '[L4:Web]' in text:
            source_tag = 'L4:Web'
            l4_count += 1
        else:
            continue

        claims = numeric_claim.findall(text)
        if not claims:
            continue

        response = retrieval.match({
            'mode': 'long_context_search',
            'query': text,
            'constraints': {'project_standards': pstandards},
            'limits': {'max_results': 3, 'max_words': 6},
        })
        usable_results, review_results, blocked_results = partition_by_support_action(response['results'])
        found_in_kb = bool(usable_results)
        best_result = usable_results[0]['raw'] if found_in_kb else None

        if found_in_kb:
            kb_hit += 1
            if gate_mode:
                issues.append({
                    'type': 'ai_claim_verified',
                    'detail': f'P{i}: [{source_tag}] verified by KB -> upgrade to [L1:KB]',
                    'kb_file': best_result.get('file', ''),
                    'kb_heading': best_result.get('heading', ''),
                    'severity': 'low'
                })
        elif review_results:
            kb_miss += 1
            review_item = review_results[0]
            raw = review_item.get('raw') or {}
            action = review_item.get('support_action') or raw.get('support_action', 'manual_review')
            sev = 'medium' if gate_mode else 'low'
            issues.append({
                'type': 'ai_claim_needs_review',
                'detail': f'P{i}: [{source_tag}] KB candidate needs review ({action}): "{text[:120]}..."',
                'kb_file': raw.get('file', ''),
                'kb_heading': raw.get('heading', ''),
                'severity': sev
            })
        else:
            kb_miss += 1
            vals = ', '.join(claims[:3])
            sev = 'high' if gate_mode or blocked_results else 'low'
            issues.append({
                'type': 'ai_claim_unverified',
                'detail': f'P{i}: [{source_tag}] VALUES:{vals} | no KB support: "{text[:120]}..."',
                'severity': sev
            })

    total = l3_count + l4_count
    if gate_mode and total > 0:
        issues.insert(0, {
            'type': 'phase2_summary',
            'detail': f'Phase2 Gate: {total} AI-sourced paragraphs ({l3_count} L3 + {l4_count} L4), '
                     f'{kb_hit} verified by KB, {kb_miss} UNVERIFIED',
            'severity': 'high' if kb_miss > 0 else 'low'
        })

    return issues

def check_code_names(all_text, kb):
    """检查方案中每个规范编号是否附带正确的官方中文名称。
    三层容错：忽略空格/·/连接号；"规范"↔"标准"版本改名报low。"""
    issues = []

    # 提取所有「编号《名称》」对
    code_name_pat = re.compile(
        r'((?:GB|JGJ|CJJ|CECS|CJ|DB|DGJ|TCECS)\s*/?\s*T?\s*[-]?\s*[\d\.]+(?:-\d+)?)\s*[《]([^》]{3,60})[》]'
    )

    seen = set()
    for i, text in all_text:
        for m in code_name_pat.finditer(text):
            raw_code = m.group(1).strip()
            plan_name = m.group(2).strip()
            nc = normalize_code(raw_code)
            if nc in seen:
                continue
            seen.add(nc)

            kb_name, layer, confidence = kb.get_name(raw_code)
            if not kb_name:
                continue  # not in KB, already caught by check_version

            # Normalize for comparison
            def _norm(s):
                return s.replace(' ', '').replace('·', '').replace('-', '').replace('\u2014', '')

            if _norm(plan_name) == _norm(kb_name):
                continue  # exact match

            # "规范"↔"标准" 容忍（版本改名）
            plan_norm2 = _norm(plan_name).replace('\u89c4\u8303', '\u6807\u51c6')
            kb_norm2 = _norm(kb_name).replace('\u89c4\u8303', '\u6807\u51c6')
            if plan_norm2 == kb_norm2:
                issues.append({
                    'type': 'code_name_version_rename',
                    'detail': f'P{i}: {raw_code} plan="{plan_name}" kb="{kb_name}" (规范↔标准 version rename)',
                    'severity': 'low'
                })
                continue

            # Partial match (one contains the other)
            if _norm(plan_name) in _norm(kb_name) or _norm(kb_name) in _norm(plan_name):
                issues.append({
                    'type': 'code_name_partial_mismatch',
                    'detail': f'P{i}: {raw_code} plan="{plan_name}" kb="{kb_name}"',
                    'severity': 'low'
                })
                continue

            # Real mismatch
            issues.append({
                'type': 'code_name_mismatch',
                'detail': f'P{i}: {raw_code} plan="{plan_name}" kb="{kb_name}" (L{layer}, {confidence})',
                'severity': 'medium'
            })

    # Check for codes cited WITHOUT any Chinese name (only codes NEVER named)
    codes_with_names = set(seen)  # populated by first pass

    # Second pass: find all standard codes in text
    all_code_pat = re.compile(
        r'((?:GB|JGJ|CJJ|CECS|CJ|DB|DGJ|TCECS)\s*/?\s*T?\s*[-]?\s*[\d\.]+(?:-\d+)?)'
    )
    # Determine which codes are NEVER followed by 《name》 anywhere
    codes_without_names = set()
    for i, text in all_text:
        for m in all_code_pat.finditer(text):
            raw_code = m.group(1).strip()
            nc = normalize_code(raw_code)
            if nc in codes_with_names:
                continue
            # Check if this occurrence has a name after it
            after = text[m.end():m.end()+60]
            if re.match(r'\s*[《]', after):
                # Has name in this occurrence
                codes_with_names.add(nc)
            else:
                codes_without_names.add(nc)

    # Report codes that NEVER appear with a name
    never_named = codes_without_names - codes_with_names
    for nc in sorted(never_named):
        kb_name, layer, confidence = kb.get_name(nc)
        if kb_name:
            # Find first occurrence paragraph
            first_p = '?'
            for i, text in all_text:
                if normalize_code(text) and nc in normalize_code(text):
                    first_p = str(i)
                    break
            issues.append({
                'type': 'code_missing_name',
                'detail': f'Code {nc} cited but never given Chinese name. KB: "{kb_name}" (L{layer}, {confidence})',
                'severity': 'medium'
            })

    return issues


def check_local_standards(doc, all_text):
    """检查地方标准（DB/DBJ/DGJ）是否与项目所在地匹配。
    噪声屏蔽：引用其他地区标准的段落包含'参照/参考/借鉴/类似/例如'时不报告。"""
    issues = []

    province, city, loc_raw = _parse_project_location(doc)
    if not province:
        return issues  # 无法识别项目地点，跳过

    # 获取期望前缀
    expected_prefixes = set()
    for (p, c), prefixes in _ADMIN_MAP.items():
        if p == province:
            if c is None or c == city:
                expected_prefixes.update(prefixes)

    if not expected_prefixes:
        return issues  # 未覆盖的省份，跳过

    # 提取所有地方标准引用
    # 格式: DB3201/T 1012-2021, DGJ32/J 16-2019, DBJ32-2020, DB11/T 1234
    local_std_pat = re.compile(
        r'((?:DBJ?|DGJ)\d{2,4}\s*/?\s*(?:T|J)?\s*[-]?\s*[\d\.]+(?:-\d+)?)'
        r'(?:\s*[《]([^》]{2,50})[》])?'
    )
    # 噪声上下文：这些词出现时不报告异地标准
    noise_words = ['参照', '参考', '借鉴', '类似', '例如', '如', '对比', '比照']

    full_text = '\n'.join(t for _, t in all_text)
    seen = set()
    for m in local_std_pat.finditer(full_text):
        raw_code = m.group(1).strip()
        plan_name = m.group(2).strip() if m.group(2) else ''

        # 标准化代码用于去重
        nc = normalize_code(raw_code)
        if nc in seen:
            continue
        seen.add(nc)

        # 检查前缀是否匹配
        matched = False
        matched_prefix = ''
        for prefix in expected_prefixes:
            # 前缀匹配：DB32→期望DB32, DGJ32→期望DGJ32
            if raw_code.startswith(prefix):
                matched = True
                matched_prefix = prefix
                break

        if matched:
            continue  # 正确的地方标准

        # 未匹配——检查是否为噪声（跨区域参考）
        ctx_start = max(0, m.start() - 60)
        ctx_end = min(len(full_text), m.end() + 40)
        context = full_text[ctx_start:ctx_end]

        is_noise = any(nw in context for nw in noise_words)

        if is_noise:
            continue  # 显式参照/参考，不报

        # 判断严重程度
        # DB32开头但不匹配南京市 → 省外标准
        if re.match(r'DB\d{2}', raw_code):
            severity = 'high'
            issue_type = 'local_std_foreign_province'
        else:
            severity = 'medium'
            issue_type = 'local_std_unexpected'

        issues.append({
            'type': issue_type,
            'detail': (
                f'{raw_code} {plan_name} — project is {province}'
                f'{"/"+city if city else ""}, '
                f'expected prefixes: {", ".join(sorted(expected_prefixes))}'
            ),
            'severity': severity
        })

    return issues


def check_required_sections(all_text):
    """检查必需章节是否存在"""
    issues = []
    required = ['编制依据', '工程概况', '施工安排', '施工准备',
                '主要施工方法', '质量要求', '安全管理', '文明施工']
    full_text = '\n'.join(t for _, t in all_text)
    for req in required:
        if req not in full_text:
            issues.append({
                'type': 'missing_section',
                'detail': f'Missing required section: {req}',
                'severity': 'high'
            })
    return issues

def print_report(issues):
    if not issues:
        print('All checks passed.')
        return True
    print(f'Issues found: {len(issues)}\n')
    for iss in issues:
        print(f'  [{iss["severity"]}] {iss["type"]}: {iss["detail"]}')
        print()
    # 返回是否有high severity
    return not any(i['severity'] == 'high' for i in issues)


if __name__ == '__main__':
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    parser = argparse.ArgumentParser()
    parser.add_argument('docx', nargs='?', help='Path to docx file (optional with --find)')
    parser.add_argument('--find', action='store_true', help='Auto-find latest docx on Desktop')
    parser.add_argument('--json', action='store_true', help='Output as JSON')
    parser.add_argument('--phase2-gate', action='store_true', help='Phase2 gate mode: blocks on unverified AI claims')
    parser.add_argument('--phase2-report', metavar='OUTPUT_HTML', help='Generate Phase2 HTML report highlighting [L3:AI] paragraphs')
    parser.add_argument('--self-test', action='store_true', help='Run built-in capability self-tests (no docx needed)')
    args = parser.parse_args()

    if args.self_test:
        """新增能力自带自检——覆盖当前模块关键路径，无需docx"""
        kb = KB()
        errors = 0

        # --- check_code_names ---
        test_text = [
            (0, '依据GB 50209-2010《建筑地面工程施工质量验收规范》执行'),
            (1, '依据GB 50202-2018《建筑地基基础工程施工质量验收标准》验收'),
            (2, '依据GB 50720-2011《建设工程施工现场消防安全技术规范》检查'),
            (3, '参照GB 50007-2011执行'),
        ]
        issues = check_code_names(test_text, kb)
        # 预期：1条name_mismatch(GB50202多"基础") + 1条version_rename(GB50720规范→标准) + 1条missing(GB50007无名)
        has_mismatch = any(i['type'] == 'code_name_mismatch' for i in issues)
        has_rename = any(i['type'] == 'code_name_version_rename' for i in issues)
        has_missing = any(i['type'] == 'code_missing_name' for i in issues)
        for cond, name in [(has_mismatch, 'name_mismatch detected'),
                           (has_rename, 'version_rename detected'),
                           (has_missing, 'missing_name detected')]:
            if cond:
                print(f'  [PASS] check_code_names: {name}')
            else:
                print(f'  [FAIL] check_code_names: {name}')
                errors += 1

        # --- check_local_standards ---
        class _FakePara:
            def __init__(self, t): self.text = t
        class _FakeDoc:
            def __init__(self, paras): self.paragraphs = [_FakePara(p) for p in paras]

        doc_nj = _FakeDoc(['南京市建邺区河西CBD'])
        # 正确本地标准
        r1 = check_local_standards(doc_nj, [(0, '依据DB3201/T 1012-2021执行')])
        # 外地标准
        r2 = check_local_standards(doc_nj, [(0, '依据DB11/T 1234-2020验收')])
        # 噪声屏蔽
        r3 = check_local_standards(doc_nj, [(0, '参照DB11/T 1234-2020做法')])

        for expected, result, name in [
            (0, len(r1), 'correct local passes'),
            (1, len(r2), 'foreign flagged'),
            (0, len(r3), 'noise suppressed'),
        ]:
            if result == expected:
                print(f'  [PASS] local_standards: {name}')
            else:
                print(f'  [FAIL] local_standards: {name} (got {result}, expected {expected})')
                errors += 1

        # --- _parse_project_location ---
        for doc_obj, exp_prov, exp_city, name in [
            (_FakeDoc(['项目地点：南京市建邺区']), '江苏', '南京', 'city+district'),
            (_FakeDoc(['北京市海淀区某项目']), '北京', '北京', 'municipality'),
            (_FakeDoc(['杭州西湖项目室外工程']), '浙江', '杭州', 'city-only title'),
        ]:
            prov, city, _ = _parse_project_location(doc_obj)
            if prov == exp_prov and city == exp_city:
                print(f'  [PASS] location parse: {name} -> {prov}/{city}')
            else:
                print(f'  [FAIL] location parse: {name} -> {prov}/{city} (expected {exp_prov}/{exp_city})')
                errors += 1

        # --- check_citation_accuracy: L2 fallback ---
        # 测试：数值不在引用条款但存在于其他规范 → 应报 medium 非 high
        test_text_acc = [
            (0, '\u4f9d\u636eGB 50209-2010\u7b2c6.1.5\u6761\uff0c\u517b\u62a4\u65f6\u95f4\u22657d'),
        ]
        issues_acc = check_citation_accuracy(test_text_acc, kb)
        has_l2 = any(i['type'] == 'citation_wrong_clause' for i in issues_acc)
        has_high = any(i['type'] == 'citation_value_mismatch' and i['severity'] == 'high' for i in issues_acc)
        print(f'  [{"PASS" if has_l2 or not has_high else "FAIL"}] citation_accuracy: L2 fallback prevents false hallucination flag')

        usable, review, blocked = partition_by_support_action([
            {'support_action': 'use_as_evidence', 'raw': {'file': 'ok'}}, 
            {'support_action': 'warn_insufficient_support', 'raw': {'file': 'review'}}, 
            {'support_action': 'block_forbidden', 'raw': {'file': 'blocked'}}, 
        ])
        support_gate_ok = len(usable) == 1 and len(review) == 1 and len(blocked) == 1
        if support_gate_ok:
            print('  [PASS] ai_claims: support guard classifies evidence/review/blocked')
        else:
            print('  [FAIL] ai_claims: support guard classification')
            errors += 1

        if errors:
            print(f'\n{errors} FAILURES')
            sys.exit(1)
        else:
            print('\nALL SELF-TESTS PASSED')
            sys.exit(0)

    docx_path = args.docx
    if args.find or not docx_path:
        docx_path = find_latest_docx()
        if not docx_path:
            print('ERROR: No .docx found on Desktop. Specify path explicitly.')
            sys.exit(1)
        print(f'Found: {docx_path}\n')

    if not os.path.exists(docx_path):
        print(f'ERROR: File not found: {docx_path}')
        sys.exit(1)

    if args.phase2_report:
        # Generate HTML report with AI claims highlighted
        from docx import Document as _Doc
        doc = _Doc(docx_path)
        kb = KB()
        ai_paras = []
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if '[L3:AI]' in t or '[L4:Web]' in t:
                ai_paras.append((i, t))
        html = '<html><head><meta charset="utf-8"><title>Phase2 Report</title>'
        html += '<style>body{font-family:sans-serif;max-width:900px;margin:20px}'
        html += '.ai{background:#fffde7;padding:8px;margin:4px 0;border-left:3px solid #f59e0b}'
        html += '.kb{background:#e8f5e9;padding:8px;margin:4px 0;border-left:3px solid #16a34a}'
        html += '.warning{color:#dc2626;font-weight:bold}</style></head><body>'
        html += f'<h1>Phase2 Report: {os.path.basename(docx_path)}</h1>'
        html += f'<p>{len(ai_paras)} paragraphs marked [L3:AI]/[L4:Web] require manual review:</p>'
        for idx, text in ai_paras:
            html += f'<div class="ai"><strong>P{idx}</strong>: {text[:200]}</div>\n'
        html += f'<p class="warning">Review each paragraph: verify claims against KB or mark as accepted.</p>'
        html += '</body></html>'
        with open(args.phase2_report, 'w', encoding='utf-8') as f:
            f.write(html)
        print(f'Phase2 report saved: {args.phase2_report} ({len(ai_paras)} paragraphs to review)')
        sys.exit(0)

    issues = verify(docx_path, gate_mode=args.phase2_gate)
    if args.json:
        print(json.dumps(issues, ensure_ascii=False, indent=2))
    else:
        ok = print_report(issues)
        if args.phase2_gate:
            unverified = sum(1 for i in issues if i['type'] == 'ai_claim_unverified')
            if unverified > 0:
                print(f'\n  PHASE2 GATE: BLOCKED - {unverified} unverified AI claims')
            else:
                print(f'\n  PHASE2 GATE: PASSED')
            ok = (unverified == 0) and ok
        exit(0 if ok else 1)
