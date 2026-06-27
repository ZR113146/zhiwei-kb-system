# -*- coding: utf-8 -*-
"""施工方案渲染引擎 —— 从content/*.md读取内容，生成docx"""
import os, re, json, tempfile, shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

# KB 图片根目录 — 用于解析相对路径
KB_IMAGES_DIR = os.path.join(os.path.dirname(__file__), '..', 'data', 'images')

class PlanRenderer:
    def __init__(self, content_dir, output_path, template_name=None):
        self.content_dir = content_dir
        self.output_path = output_path
        self.template = self._load_template(template_name)
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    def _load_template(self, template_name):
        """Load format template from templates/*.json. Fallback to built-in defaults."""
        tpl_dir = os.path.join(os.path.dirname(__file__), '..', 'templates')
        name = template_name or 'default'
        tpl_path = os.path.join(tpl_dir, f'{name}.json')
        if os.path.exists(tpl_path):
            with open(tpl_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        if template_name:
            print(f'  WARNING: Template "{template_name}" not found at {tpl_path} — using built-in defaults')
        return None  # signal fallback

    def _setup_page(self):
        t = self.template.get('page', {}) if self.template else {}
        for s in self.doc.sections:
            s.page_width = Cm(t.get('width_cm', 21.0))
            s.page_height = Cm(t.get('height_cm', 29.7))
            s.top_margin = Cm(t.get('top_margin_cm', 2.5))
            s.bottom_margin = Cm(t.get('bottom_margin_cm', 2.0))
            s.left_margin = Cm(t.get('left_margin_cm', 2.8))
            s.right_margin = Cm(t.get('right_margin_cm', 2.8))

    def _setup_styles(self):
        t = self.template.get('body', {}) if self.template else {}
        font_name = t.get('font', '\u5b8b\u4f53')
        font_size = t.get('size_pt', 12)
        line_sp = t.get('line_spacing', 1.5)
        indent = t.get('first_line_indent_pt', 24)

        ns = self.doc.styles['Normal']
        ns.font.name = font_name
        ns.font.size = Pt(font_size)
        ns.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        ns.paragraph_format.line_spacing = line_sp
        ns.paragraph_format.first_line_indent = Pt(indent)

    # ---- 标题 ----
    def h1(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(30)
        p.paragraph_format.space_after = Pt(18)
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(text); r.font.size = Pt(16); r.font.bold = True
        r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def h2(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(22)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(text); r.font.size = Pt(14); r.font.bold = True

    def h3(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(text); r.font.size = Pt(14); r.font.bold = True

    def h4(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(text); r.font.size = Pt(12); r.font.bold = True

    # ---- 正文 ----
    def p(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(text)

    def li(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(text)

    def blank(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)

    # ---- 表格 ----
    def _set_cell_shading(self, cell, hex_fill):
        """给单元格设置底纹填充色 (hex, 如 'D9D9D9')。"""
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_fill)
        tcPr.append(shd)

    def _set_repeat_header(self, row):
        """标记表头行，使其在跨页时自动重复。"""
        from docx.oxml import OxmlElement
        trPr = row._tr.get_or_add_trPr()
        th = OxmlElement('w:tblHeader')
        th.set(qn('w:val'), 'true')
        trPr.append(th)

    def _set_fixed_layout(self, table):
        """固定列宽布局，避免 Word 自动撑列导致列宽错乱。"""
        from docx.oxml import OxmlElement
        tblPr = table._tbl.tblPr
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)

    def _style_cell(self, cell, text, *, bold=False, font='宋体', size=10,
                    align=WD_ALIGN_PARAGRAPH.CENTER, fill=None):
        """统一单元格样式：清空→写入→字体/字号/加粗/对齐/垂直居中/底纹。"""
        cell.text = ''
        para = cell.paragraphs[0]
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.alignment = align
        r = para.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.name = font
        r.element.rPr.rFonts.set(qn('w:eastAsia'), font)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if fill:
            self._set_cell_shading(cell, fill)

    def table(self, headers, rows):
        ncols, nrows = len(headers), len(rows) + 1
        t = self.doc.add_table(rows=nrows, cols=ncols, style='Table Grid')
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        self._set_fixed_layout(t)

        # 页面可用宽度: A4(21cm) - 左右边距(2.8×2) = 15.4cm
        PAGE_WIDTH = 15.4

        # Calculate column widths: scan ALL cells, count CJK chars separately
        col_weights = []
        for ci in range(ncols):
            all_texts = [str(headers[ci])]
            for row in rows:
                if ci < len(row):
                    all_texts.append(str(row[ci]))
            # Find the max-width cell for this column
            max_effective = 0
            for txt in all_texts:
                cjk = sum(1 for c in txt if ord(c) > 127)
                ascii_count = len(txt) - cjk
                effective = ascii_count * 0.5 + cjk * 1.0  # CJK ≈ 2x ASCII width
                max_effective = max(max_effective, effective)
            col_weights.append(max(max_effective, 2.0))  # minimum weight

        total_weight = sum(col_weights)
        col_widths = [Cm(PAGE_WIDTH * w / total_weight) for w in col_weights]

        # 列宽需同时设置到 table.columns 和每个单元格，固定布局才生效
        for ci in range(ncols):
            if ci < len(col_widths):
                t.columns[ci].width = col_widths[ci]

        # Header row — 灰底 + 黑体加粗，并标记跨页重复
        for ci, h in enumerate(headers):
            c = t.rows[0].cells[ci]
            self._style_cell(c, str(h), bold=True, font='黑体', size=10, fill='D9D9D9')
            if ci < len(col_widths):
                c.width = col_widths[ci]
        self._set_repeat_header(t.rows[0])

        # Data rows — 宋体正文
        for ri, row in enumerate(rows):
            if len(row) > ncols:
                print(f'  WARNING: Table row {ri+1} has {len(row)} cols (header: {ncols}) — extra columns dropped')
            for ci, val in enumerate(row):
                if ci < ncols:
                    c = t.rows[ri+1].cells[ci]
                    self._style_cell(c, str(val).strip(), bold=False, font='宋体', size=10)
                    if ci < len(col_widths):
                        c.width = col_widths[ci]
        self.blank()

    # ---- 图片 ----
    def img(self, path, width_cm=14):
        # 解析路径: 相对路径 → 先查 content_dir, 再查 KB_IMAGES_DIR
        resolved = path
        if not os.path.isabs(path):
            # 1. 相对于 content_dir
            candidate = os.path.join(self.content_dir, path)
            if os.path.exists(candidate):
                resolved = candidate
            # 2. KB 图片目录 (如 images/xxx.jpg)
            elif path.startswith('images/') or path.startswith('images\\'):
                candidate2 = os.path.join(KB_IMAGES_DIR, os.path.basename(path))
                if os.path.exists(candidate2):
                    resolved = candidate2
                else:
                    # 用哈希文件名查找
                    candidate3 = os.path.join(KB_IMAGES_DIR, path.replace('images/', '').replace('images\\', ''))
                    if os.path.exists(candidate3):
                        resolved = candidate3
        if os.path.exists(resolved):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            try:
                p.add_run().add_picture(resolved, width=Cm(width_cm))
            except Exception:
                self.img_placeholder(os.path.basename(path))
        else:
            # 图片文件不存在 → 占位符
            self.img_placeholder(os.path.basename(path))

    def img_placeholder(self, description):
        """图片占位符 - 用户手动替换"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(16)
        r = p.add_run(f'【此处插入{description}】')
        r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ---- 封面 ----
    def cover(self, title, subtitle, company, date, reviewer=None):
        for _ in range(6): self.blank()
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(title); r.font.size = Pt(22); r.font.bold = True
        r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(subtitle); r.font.size = Pt(22); r.font.bold = True
        r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        for _ in range(6): self.blank()
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Pt(0)
        p.add_run(f'编制单位：{company}').font.size = Pt(14)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Pt(0)
        p.add_run(f'编制日期：{date}').font.size = Pt(14)
        if reviewer:
            self.blank()
            p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Pt(0)
            p.add_run(f'审批单位：{reviewer}').font.size = Pt(12)
            p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Pt(0)
            p.add_run(f'审批日期：{date}').font.size = Pt(12)
        self.doc.add_page_break()

    # ---- 解析内容文件 ----
    def render_content_file(self, filename):
        """解析content/*.md文件并渲染"""
        path = os.path.join(self.content_dir, filename)
        if not os.path.exists(path):
            print(f'WARNING: {path} not found, skipping')
            return

        with open(path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        i = 0
        in_code_block = False
        while i < len(lines):
            line = lines[i].rstrip()

            # 代码块 (跳过围栏行, 内容渲染为等宽小字)
            if line.startswith('```'):
                in_code_block = not in_code_block
                i += 1
                continue
            if in_code_block:
                if line:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.left_indent = Cm(1.0)
                    r = p.add_run(line)
                    r.font.size = Pt(9)
                    r.font.name = 'Consolas'
                i += 1
                continue

            # LaTeX 数学公式: $$...$$ → 居中, $...$ → 内联保留
            if line.startswith('$$') and line.endswith('$$'):
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                r = p.add_run(line.strip('$ '))
                r.font.size = Pt(10)
                i += 1
                continue

            # 空行
            if not line:
                i += 1
                continue

            # 标题
            if line.startswith('# '):
                self.h1(line[2:])
            elif line.startswith('## '):
                self.h2(line[3:])
            elif line.startswith('### '):
                self.h3(line[4:])
            elif line.startswith('#### '):
                self.h4(line[5:])

            # 列表项
            elif line.startswith('- '):
                self.li(line[2:])

            # 表格
            elif line.startswith('|'):
                rows = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row = [c.strip() for c in lines[i].split('|')[1:-1]]
                    rows.append(row)
                    i += 1
                if rows:
                    # Skip separator row
                    data_rows = [r for r in rows if not all(c.replace('-','').replace(':','').strip()=='' for c in r)]
                    if len(data_rows) > 1:
                        self.table(data_rows[0], data_rows[1:])
                continue

            # 图片占位符
            elif line.startswith('![IMG]'):
                desc = line.replace('![IMG]', '').strip('() ')
                self.img_placeholder(desc)

            # 图片文件
            elif line.startswith('![') and '](' in line:
                path_part = line.split('](')[1].rstrip(')')
                self.img(path_part)

            # 分页
            elif line == '---':
                self.doc.add_page_break()

            # <details> HTML 标签: 提取 summary 后跳过 (图片由 Mermaid 提取处理)
            elif line.startswith('<details>') or line.startswith('</details>'):
                i += 1
                continue
            elif line.startswith('<summary>'):
                i += 1
                continue

            # 普通段落
            elif line:
                self.p(line)

            i += 1

    def load_project_config(self):
        """Load project info from content/project.json, fallback to defaults"""
        config_path = os.path.join(self.content_dir, 'project.json')
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {
            'title': '【项目名称】',
            'subtitle': '专项施工方案',
            'company': '【编制单位】',
            'date': '【日期】',
            'reviewer': None
        }

    def _extract_mermaid_blocks(self):
        """预扫描所有 content/*.md，提取 ```mermaid 代码块为 .mmd 文件，替换为占位符"""
        md_files = sorted([f for f in os.listdir(self.content_dir) if f.endswith('.md')])
        extracted = 0
        for mdf in md_files:
            path = os.path.join(self.content_dir, mdf)
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            # 查找 ```mermaid ... ``` 块
            pattern = re.compile(r'```mermaid\s*\n(.*?)```', re.DOTALL)
            blocks = pattern.findall(content)
            if not blocks:
                continue
            modified = content
            for idx, block in enumerate(blocks):
                block = block.strip()
                if not block:
                    continue
                # 生成唯一文件名
                basename = os.path.splitext(mdf)[0]
                mmd_name = f'{basename}_flow{idx+1}.mmd'
                mmd_path = os.path.join(self.content_dir, mmd_name)
                with open(mmd_path, 'w', encoding='utf-8') as f:
                    f.write(block)
                # 替换为占位符
                old = f'```mermaid\n{block}\n```'
                new = f'![IMG]({os.path.splitext(mmd_name)[0]})'
                modified = modified.replace(old, new)
                extracted += 1
            if modified != content:
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(modified)
        if extracted:
            print(f'  Mermaid: extracted {extracted} block(s) from content files')

    def build(self):
        """构建完整方案"""
        cfg = self.load_project_config()

        # 预提取 mermaid 代码块
        self._extract_mermaid_blocks()

        # 封面
        self.cover(
            title=cfg.get('title', ''),
            subtitle=cfg.get('subtitle', '专项施工方案'),
            company=cfg.get('company', ''),
            date=cfg.get('date', ''),
            reviewer=cfg.get('reviewer')
        )

        # 逐章渲染
        for ch in ['ch01','ch02','ch03','ch04','ch05','ch06','ch07','ch08','ch09']:
            self.render_content_file(f'{ch}.md')

        # 自动生成并嵌入流程图（如果Mermaid源文件存在）
        self._embed_flowcharts()

        self.doc.save(self.output_path)
        print(f'Generated: {self.output_path} ({os.path.getsize(self.output_path)} bytes)')

    def _write_mermaid_assets(self):
        """生成 mermaid-cli 的渲染配置与 CSS，返回 (config_path, css_path)。
        关键点：中文字体、flowchart 直角连线(step)、加大节点间距与内边距，
        让流程图更贴近工程文档观感，且高 DPI 渲染不发虚。
        允许模板通过 flowchart.font / flowchart.curve 覆盖默认值。"""
        fc = (self.template.get('flowchart', {}) if self.template else {})
        font = fc.get('font', 'Microsoft YaHei, SimHei, sans-serif')
        curve = fc.get('curve', 'step')  # step=直角折线，工程图首选
        config = {
            'theme': 'base',
            'themeVariables': {
                'fontFamily': font,
                'fontSize': '16px',
                'lineColor': '#333333',
                'primaryColor': '#ffffff',
                'primaryBorderColor': '#333333',
                'primaryTextColor': '#000000',
            },
            'flowchart': {
                'curve': curve,
                'nodeSpacing': 60,
                'rankSpacing': 70,
                'padding': 16,
                'htmlLabels': True,
                'useMaxWidth': False,
            },
        }
        css = (
            f'.node rect, .node polygon, .node circle, .node path '
            f'{{ stroke-width: 1.5px; }} '
            f'.edgeLabel {{ font-family: {font}; background: #ffffff; }} '
            f'.nodeLabel, .edgeLabel, text {{ font-family: {font}; }}'
        )
        cfg_path = os.path.join(self.content_dir, '.mermaid_config.json')
        css_path = os.path.join(self.content_dir, '.mermaid_theme.css')
        with open(cfg_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        with open(css_path, 'w', encoding='utf-8') as f:
            f.write(css)
        return cfg_path, css_path

    def _embed_flowcharts(self):
        """查找content/*.mmd并自动渲染嵌入到对应位置"""
        import glob as _glob
        mmd_files = sorted(_glob.glob(os.path.join(self.content_dir, '*.mmd')))
        if not mmd_files:
            return

        mmdc_cfg = (self.template.get('flowchart', {}).get('mmdc_js', '') if self.template else '')
        mmdc_js = os.path.expanduser(mmdc_cfg) if mmdc_cfg else os.path.join(
            os.path.expanduser('~'), 'AppData', 'Roaming', 'npm',
            'node_modules', '@mermaid-js', 'mermaid-cli', 'src', 'cli.js')

        # 渲染质量配置 — 中文字体、直角连线、加大留白，输出高清 PNG
        cfg_path, css_path = self._write_mermaid_assets()

        rendered_png = 0
        embedded_count = 0
        for mmd_path in mmd_files:
            png_path = mmd_path.replace('.mmd', '.png')
            # Render if PNG doesn't exist or MMD is newer
            if not os.path.exists(png_path) or os.path.getmtime(mmd_path) > os.path.getmtime(png_path):
                import subprocess
                try:
                    subprocess.run(['node', mmdc_js, '-i', mmd_path, '-o', png_path,
                                   '-b', 'white', '-s', '3',
                                   '-c', cfg_path, '-C', css_path],
                                   capture_output=True, text=True,
                                   encoding='utf-8', errors='replace', timeout=30)
                    rendered_png += 1
                except subprocess.TimeoutExpired:
                    print(f'  WARNING: mermaid-cli timed out for {os.path.basename(mmd_path)}')
                except FileNotFoundError:
                    print(f'  WARNING: mermaid-cli not found — install with "npm i -g @mermaid-js/mermaid-cli"')
                    return  # No point continuing without mermaid-cli
                except Exception as e:
                    print(f'  WARNING: mermaid-cli failed for {os.path.basename(mmd_path)}: {e}')

            # Find placeholder text from filename and insert image
            if os.path.exists(png_path):
                fname = os.path.basename(mmd_path).replace('.mmd', '')
                for i, p in enumerate(self.doc.paragraphs):
                    # Only match IMG placeholder paragraphs (contain both fname AND placeholder brackets)
                    if fname in p.text and '\u3010' in p.text:
                        # Remove ALL existing content (runs + drawings)
                        for child in list(p._element):
                            p._element.remove(child)
                        # Add image — height-first: cap at 22cm
                        from PIL import Image as PILImage
                        try:
                            img = PILImage.open(png_path)
                            pw, ph = img.size
                            MAX_H = 22.0
                            if ph / pw * 14 > MAX_H:
                                p.add_run().add_picture(png_path, height=Cm(MAX_H))
                            else:
                                p.add_run().add_picture(png_path, width=Cm(14))
                            embedded_count += 1
                        except Exception:
                            try:
                                p.add_run().add_picture(png_path, width=Cm(14))
                                embedded_count += 1
                            except Exception as e2:
                                print(f'  WARNING: Failed to insert {fname}: {e2}')
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        break

        # Post-verification: check for remaining flowchart placeholders
        if embedded_count or rendered_png:
            remaining = 0
            for p in self.doc.paragraphs:
                if '\u6d41\u7a0b\u56fe' in p.text and '\u3010' in p.text:
                    remaining += 1
            if remaining > 0:
                print(f'  WARNING: {remaining} flowchart placeholder(s) still present')
        if rendered_png or embedded_count:
            print(f'  Flowcharts: {rendered_png} rendered, {embedded_count} embedded')


# ====== 独立运行 ======
if __name__ == '__main__':
    import sys
    content_dir = sys.argv[1] if len(sys.argv) > 1 else 'content'
    output = sys.argv[2] if len(sys.argv) > 2 else 'output.docx'
    PlanRenderer(content_dir, output).build()
