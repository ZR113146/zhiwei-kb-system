"""Evaluate plan-writer citation integration against KB structured citations."""

import json
import os
import sys
from datetime import datetime

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
for rel in ("", "kb_core", "plan_writer"):
    path = os.path.join(ROOT, rel) if rel else ROOT
    if path not in sys.path:
        sys.path.insert(0, path)

from kb_core.kb import KB  # noqa: E402
from docx import Document  # noqa: E402
from kb_auditor import (  # noqa: E402
    _resolve_citation,
    append_citation_audit_summary_to_docx,
    resolve_for_chapter,
    write_citation_audit_summary,
)

DEFAULT_JSON = os.path.join(ROOT, "eval", "plan_writer_citation_eval_report.json")
DEFAULT_MD = os.path.join(ROOT, "eval", "plan_writer_citation_eval_report.md")


def run_checks():
    kb = KB()
    rows = []

    chapter_rows = resolve_for_chapter("ch05", "paving", kb)
    has_citations = any(item.get("citations") for item in chapter_rows)
    rows.append({
        "name": "resolve_for_chapter_returns_citations",
        "ok": has_citations,
        "detail": f"rows={len(chapter_rows)} citations={sum(len(item.get('citations', [])) for item in chapter_rows)}",
    })

    resolved = _resolve_citation(kb, "GB50202-2018", "5.1.3", "50 ~m", {"context": "灌注桩混凝土强度检验"})
    citation = resolved.get("citation") or {}
    rows.append({
        "name": "resolve_citation_attaches_kb_citation",
        "ok": bool(citation) and citation.get("audit_status") == "pass" and citation.get("clause_type") == "normative",
        "detail": json.dumps({
            "in_clause": resolved.get("in_clause"),
            "clause_exists": resolved.get("clause_exists"),
            "audit_status": citation.get("audit_status"),
            "clause_type": citation.get("clause_type"),
            "source_file": citation.get("source_file"),
        }, ensure_ascii=False),
    })

    summary_base = os.path.join(ROOT, "eval", "plan_writer_citation_summary_sample")
    summary_written = write_citation_audit_summary("sample.docx", [{
        "para": 1,
        "code": "GB50202-2018",
        "clause_ref": "5.1.3",
        "claimed": "50 ~m",
        "resolution": resolved,
        "chapter_ok": True,
    }], summary_base)
    rows.append({
        "name": "write_citation_audit_summary_outputs_files",
        "ok": os.path.exists(summary_written["json"]) and os.path.exists(summary_written["md"]),
        "detail": json.dumps({"json": summary_written["json"], "md": summary_written["md"]}, ensure_ascii=False),
    })

    sample_docx = os.path.join(ROOT, "eval", "plan_writer_citation_sample.docx")
    appended_docx = os.path.join(ROOT, "eval", "plan_writer_citation_sample_with_summary.docx")
    doc = Document()
    doc.add_paragraph("引用 GB 50202-2018 第5.1.3条，50 ~m。")
    doc.save(sample_docx)
    append_result = append_citation_audit_summary_to_docx(sample_docx, [{
        "para": 1,
        "code": "GB50202-2018",
        "clause_ref": "5.1.3",
        "claimed": "50 ~m",
        "resolution": resolved,
        "chapter_ok": True,
    }], appended_docx)
    appended = Document(append_result["docx"])
    has_heading = any("引用审计摘要" in p.text for p in appended.paragraphs)
    rows.append({
        "name": "append_citation_audit_summary_to_docx_outputs_docx",
        "ok": os.path.exists(append_result["docx"]) and has_heading,
        "detail": json.dumps({"docx": append_result["docx"], "has_heading": has_heading}, ensure_ascii=False),
    })

    return rows


def write_reports(rows, json_path=DEFAULT_JSON, md_path=DEFAULT_MD):
    summary = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "total": len(rows),
        "ok": sum(1 for row in rows if row["ok"]),
        "failures": [row for row in rows if not row["ok"]],
        "rows": rows,
    }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    lines = [
        "# Plan Writer Citation Evaluation",
        "",
        f"Generated: {summary['generated_at']}",
        f"Total: {summary['total']}",
        f"OK: {summary['ok']}",
        "",
        "## Checks",
        "",
    ]
    for row in rows:
        mark = "PASS" if row["ok"] else "FAIL"
        lines.append(f"- {mark}: {row['name']} - {row['detail']}")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")
    return summary


def main():
    summary = write_reports(run_checks())
    print(json.dumps({k: v for k, v in summary.items() if k != "rows"}, ensure_ascii=False, indent=2))
    raise SystemExit(0 if not summary["failures"] else 1)


if __name__ == "__main__":
    main()
