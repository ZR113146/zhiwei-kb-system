"""Build a manual review sheet for citation audit review items."""

import csv
import json
import os
import sys
from pathlib import Path

from docx import Document

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
for rel in ("", "kb_core"):
    path = os.path.join(ROOT, rel) if rel else ROOT
    if path not in sys.path:
        sys.path.insert(0, path)

from kb import KB  # noqa: E402


def build_review(summary_json, docx_path, out_base):
    summary = json.loads(Path(summary_json).read_text(encoding="utf-8"))
    doc = Document(docx_path)
    kb = KB()
    rows = []
    for item in summary.get("items", []):
        reasons = item.get("review_reasons") or []
        if not reasons:
            continue
        para_idx = item.get("para")
        para_text = doc.paragraphs[para_idx].text.strip() if isinstance(para_idx, int) and para_idx < len(doc.paragraphs) else ""
        code = item.get("code", "")
        clause = item.get("clause_ref", "")
        clause_data = kb.read_clause_full(code, clause) if code and clause else {}
        kb_text = (clause_data.get("clause_text") or clause_data.get("text") or "").replace("\n", " ")
        rows.append({
            "para": para_idx,
            "code": code,
            "official_code": item.get("official_code", ""),
            "clause_ref": clause,
            "claimed": item.get("claimed", ""),
            "review_reasons": ";".join(reasons),
            "audit_status": item.get("audit_status", ""),
            "version_status": item.get("version_status", ""),
            "clause_type": item.get("clause_type", ""),
            "standard_name": item.get("standard_name", ""),
            "para_text": para_text,
            "kb_clause_type": clause_data.get("clause_type", ""),
            "kb_match_method": clause_data.get("match_method", ""),
            "kb_failure_reason": clause_data.get("failure_reason", ""),
            "kb_source_file": clause_data.get("source_file", ""),
            "kb_text_preview": kb_text[:500],
            "classification": "",
            "review_note": "",
        })
    csv_path = out_base + ".csv"
    md_path = out_base + ".md"
    os.makedirs(os.path.dirname(os.path.abspath(out_base)), exist_ok=True)
    fieldnames = list(rows[0].keys()) if rows else []
    with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)
    lines = ["# Citation Review Items", ""]
    for row in rows:
        lines.extend([
            f"## P{row['para']} {row['official_code'] or row['code']} §{row['clause_ref']}",
            "",
            f"- Claimed: {row['claimed']}",
            f"- Reasons: {row['review_reasons']}",
            f"- Audit: {row['audit_status']} / {row['version_status']} / {row['clause_type']}",
            f"- KB: {row['kb_clause_type']} / {row['kb_match_method']} / {row['kb_failure_reason']}",
            f"- Source: {row['kb_source_file']}",
            "",
            "Paragraph:",
            "",
            row['para_text'],
            "",
            "KB preview:",
            "",
            row['kb_text_preview'],
            "",
        ])
    Path(md_path).write_text("\n".join(lines), encoding="utf-8")
    return {"csv": csv_path, "md": md_path, "count": len(rows)}


if __name__ == "__main__":
    if len(sys.argv) != 4:
        raise SystemExit("usage: python build_citation_review_sheet.py summary.json docx out_base")
    result = build_review(sys.argv[1], sys.argv[2], sys.argv[3])
    print(json.dumps(result, ensure_ascii=False, indent=2))
