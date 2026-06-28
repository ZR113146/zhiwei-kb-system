#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""diff_docx: compare two docx versions, output structured changes.
Used between modification rounds to detect residuals, drifts, and regressions.

Usage:
  python diff_docx.py old.docx new.docx          # compare two docx
  python diff_docx.py --find new.docx             # auto-find previous dump
  python diff_docx.py --dump new.docx             # dump + compare to last dump
"""

import os, re, sys, json, argparse, hashlib
from datetime import datetime
from docx import Document
from _utils import find_latest_docx

DUMP_DIR = os.path.join(os.path.dirname(__file__), '..', '..', '..', '..', '..', 'Desktop')
LAST_DUMP = os.path.join(DUMP_DIR, '_docx_last_dump.txt')


def dump_docx(docx_path):
    """Extract all paragraphs and tables to a structured list"""
    doc = Document(docx_path)
    entries = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            entries.append({'type': 'para', 'idx': i, 'text': t})
    for ti, t in enumerate(doc.tables):
        for ri, r in enumerate(t.rows):
            cells = [c.text.strip() for c in r.cells]
            entries.append({'type': 'table', 'tid': ti, 'rid': ri, 'cells': cells})
    return entries


def hash_text(text):
    return hashlib.md5(text.encode('utf-8')).hexdigest()[:8]


def diff_entries(old_entries, new_entries):
    """Compare two entry lists, return changes"""
    changes = []
    old_map = {hash_text(e.get('text', str(e.get('cells', '')))): e for e in old_entries if e['type'] == 'para'}
    new_map = {hash_text(e.get('text', str(e.get('cells', '')))): e for e in new_entries if e['type'] == 'para'}

    old_hashes = set(old_map.keys())
    new_hashes = set(new_map.keys())

    # Added (in new but not old)
    added = new_hashes - old_hashes
    for h in added:
        e = new_map[h]
        changes.append({'type': 'added', 'idx': e['idx'], 'text': e['text'][:200]})

    # Removed (in old but not new)
    removed = old_hashes - new_hashes
    for h in removed:
        e = old_map[h]
        changes.append({'type': 'removed', 'idx': e['idx'], 'text': e['text'][:200]})

    return changes


def find_suspicious(new_entries):
    """Detect suspicious patterns without an old version"""
    issues = []

    # 1. Old+new text coexistence (residual from incomplete replacement)
    conflict_pairs = [
        ('\u67e5\u6728', '\u82b1\u5883\u82d7'),  # 乔木 + 花境苗
        ('\u6d4e\u6728', '\u82b1\u5883'),  # 树 + 花
    ]
    for e in new_entries:
        if e['type'] != 'para':
            continue
        for old_w, new_w in conflict_pairs:
            if old_w in e['text'] and new_w in e['text']:
                issues.append({
                    'type': 'text_coexistence',
                    'idx': e['idx'],
                    'detail': f'Both "{old_w}" and "{new_w}" in same paragraph',
                    'text': e['text'][:150]
                })

    # 2. Duplicate adjacent paragraphs
    for i in range(len(new_entries) - 1):
        if new_entries[i]['type'] == 'para' and new_entries[i+1]['type'] == 'para':
            if new_entries[i]['text'] == new_entries[i+1]['text']:
                issues.append({
                    'type': 'duplicate_para',
                    'idx': new_entries[i]['idx'],
                    'detail': f'Duplicate of next paragraph',
                    'text': new_entries[i]['text'][:100]
                })

    # 3. Placeholder residue
    placeholders = ['\u25a0', '\u5f85\u8865\u5145', 'TBD', '\u6309\u56fe\u7eb8\u786e\u5b9a', '____', '...']
    for e in new_entries:
        if e['type'] != 'para':
            continue
        for ph in placeholders:
            if ph in e['text']:
                issues.append({
                    'type': 'placeholder',
                    'idx': e['idx'],
                    'detail': f'Placeholder: "{ph}"',
                    'text': e['text'][:100]
                })

    # 4. Numeric changes (flag significant value shifts)
    numeric_pattern = re.compile(r'(\d+\.?\d*)\s*(m|mm|cm|kN|kPa|d|\u5929)')
    for e in new_entries:
        if e['type'] != 'para':
            continue
        nums = numeric_pattern.findall(e['text'])
        for val, unit in nums:
            fval = float(val)
            if unit in ('m',) and fval <= 0.6:
                pass  # normal small values
            elif unit in ('kN',) and fval > 50:
                issues.append({
                    'type': 'suspicious_value',
                    'idx': e['idx'],
                    'detail': f'Large value: {val}{unit}',
                    'text': e['text'][:100]
                })

    return issues


def print_report(changes, issues, old_count, new_count):
    print(f'Paragraphs: {old_count} -> {new_count}')
    if changes:
        added_n = sum(1 for c in changes if c['type'] == 'added')
        removed_n = sum(1 for c in changes if c['type'] == 'removed')
        print(f'Changes: +{added_n} added, -{removed_n} removed')
        print()
        for c in changes:
            tag = '+' if c['type'] == 'added' else '-'
            print(f'  [{tag}] P{c["idx"]}: {c["text"][:120]}')
            print()

    if issues:
        print(f'\nSuspicious patterns: {len(issues)}')
        for iss in issues:
            print(f'  [{iss["type"]}] P{iss["idx"]}: {iss.get("detail","")}')
            print(f'    {iss["text"][:120]}')
            print()

    if not changes and not issues:
        print('No changes or issues detected.')




if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    parser = argparse.ArgumentParser()
    parser.add_argument('docx', nargs='?', help='Path to new docx file')
    parser.add_argument('--old', help='Path to old docx for comparison')
    parser.add_argument('--find', action='store_true', help='Auto-find latest docx; --old for comparison')
    parser.add_argument('--suspicious-only', action='store_true', help='Only run suspicious pattern check')
    parser.add_argument('--dump', action='store_true', help='Save dump for future comparison')
    args = parser.parse_args()

    docx_path = args.docx
    if args.find or not docx_path:
        docx_path = find_latest_docx()
        if not docx_path:
            print('ERROR: No .docx found on Desktop.')
            sys.exit(1)
        print(f'Found: {docx_path}\n')

    if not os.path.exists(docx_path):
        print(f'ERROR: File not found: {docx_path}')
        sys.exit(1)

    new_entries = dump_docx(docx_path)

    # Save dump if requested
    if args.dump:
        with open(LAST_DUMP, 'w', encoding='utf-8') as f:
            for e in new_entries:
                if e['type'] == 'para':
                    f.write(f'P{e["idx"]}|{e["text"]}\n')
                else:
                    f.write('T{}R{}|{}\n'.format(e['tid'], e['rid'], ' | '.join(e['cells'])))
        print(f'Dump saved to {LAST_DUMP}')

    if args.suspicious_only or not args.old:
        issues = find_suspicious(new_entries)
        print_report([], issues, 0, len([e for e in new_entries if e['type'] == 'para']))
    else:
        if os.path.exists(args.old):
            old_entries = dump_docx(args.old)
            changes = diff_entries(old_entries, new_entries)
            issues = find_suspicious(new_entries)
            old_para_count = len([e for e in old_entries if e['type'] == 'para'])
            new_para_count = len([e for e in new_entries if e['type'] == 'para'])
            print_report(changes, issues, old_para_count, new_para_count)
        else:
            print(f'Old file not found: {args.old}')
            issues = find_suspicious(new_entries)
            print_report([], issues, 0, len([e for e in new_entries if e['type'] == 'para']))
