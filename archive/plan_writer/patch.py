# -*- coding: utf-8 -*-
"""增量编辑: backup→定位原文→替换→保存"""
import os, shutil, json, argparse, re
from datetime import datetime
from docx import Document
from _utils import backup_docx as backup


def locate_paragraph(doc, text, context_before=None):
    """在docx中定位包含指定文本的段落（含正文+表格cell）"""
    candidates = []
    for i, p in enumerate(doc.paragraphs):
        if text in p.text:
            candidates.append((i, p, None))
    # Also search table cells
    for ti, tab in enumerate(doc.tables):
        for ri, row in enumerate(tab.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    if text in p.text:
                        candidates.append((ti, p, (ri, ci, pi)))
    if len(candidates) == 1:
        return candidates[0][1]
    elif len(candidates) > 1 and context_before:
        for idx, p, cell_ref in candidates:
            if cell_ref is None:
                prev_idx = idx - 1
                while prev_idx >= 0 and not doc.paragraphs[prev_idx].text.strip():
                    prev_idx -= 1
                if prev_idx >= 0 and context_before in doc.paragraphs[prev_idx].text:
                    return p
    elif len(candidates) == 1:
        return candidates[0][1]
    return None

def apply_single_edit(doc, edit):
    """执行单个编辑指令"""
    cmd = edit.get('action', 'replace')
    old_text = edit.get('locate', '')
    new_text = edit.get('new_text', '')
    context = edit.get('context_before')

    if cmd == 'replace':
        p = locate_paragraph(doc, old_text, context)
        if p:
            # 保留原有格式，替换文本
            for run in p.runs:
                run.text = ''
            p.runs[0].text = new_text if p.runs else None
            if not p.runs:
                p.add_run(new_text)
            return True
        else:
            print(f'  WARNING: text not found: "{old_text[:60]}..."')
            return False

    elif cmd == 'append_after':
        p = locate_paragraph(doc, old_text)
        if p:
            new_p = doc.add_paragraph()
            new_p.add_run(new_text)
            # 移动到最后
            p._element.addnext(new_p._element)
            return True

    elif cmd == 'insert_before':
        p = locate_paragraph(doc, old_text)
        if p:
            new_p = doc.add_paragraph()
            new_p.add_run(new_text)
            p._element.addprevious(new_p._element)
            return True

    return False

def apply_edits(docx_path, edits):
    """批量执行编辑指令"""
    # backup
    backup(docx_path)

    doc = Document(docx_path)
    success = 0
    for i, edit in enumerate(edits):
        ok = apply_single_edit(doc, edit)
        if ok:
            success += 1
        else:
            print(f'  Edit {i+1}/{len(edits)} failed')

    doc.save(docx_path)
    print(f'\nApplied: {success}/{len(edits)} edits')
    return success


def replace_all(docx_path, old_text, new_text):
    """Full-document search & replace — paragraphs AND table cells (no backup, caller backs up)"""
    from docx import Document
    doc = Document(docx_path)
    count = 0

    def _do_replace(para):
        nonlocal count
        if old_text in para.text:
            if para.runs:
                for r in para.runs[1:]:
                    r.text = ''
                para.runs[0].text = para.text.replace(old_text, new_text)
            else:
                para.text = para.text.replace(old_text, new_text)
            count += 1

    for p in doc.paragraphs:
        _do_replace(p)

    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _do_replace(p)

    doc.save(docx_path)
    return count

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('docx', help='Path to docx file')
    parser.add_argument('--edits', help='JSON file with edits array')
    parser.add_argument('--single', help='Single edit as JSON string')
    parser.add_argument('--replace-all', nargs=2, metavar=('OLD', 'NEW'),
                       help='Full-document search & replace (paragraphs + tables)')
    parser.add_argument('--find', action='store_true', help='Auto-find latest docx on Desktop')
    args = parser.parse_args()

    if args.find or not os.path.exists(args.docx):
        docx_path = None
        import glob as _glob
        desktop = os.path.join(os.environ.get('USERPROFILE', os.path.expanduser('~')), 'Desktop')
        candidates = []
        for f in _glob.glob(os.path.join(desktop, '*.docx')):
            if '.bak' not in f and not os.path.basename(f).startswith('~$'):
                candidates.append((os.path.getmtime(f), f))
        candidates.sort(reverse=True)
        docx_path = candidates[0][1] if candidates else None
        if not docx_path:
            print('ERROR: No .docx found on Desktop. Specify path explicitly.')
            sys.exit(1)
        if args.find:
            print('Found: {}'.format(docx_path))
        args.docx = docx_path

    if args.replace_all:
        backup(args.docx)
        count = replace_all(args.docx, args.replace_all[0], args.replace_all[1])
        print('Replaced {} instances of "{}" -> "{}"'.format(count, args.replace_all[0], args.replace_all[1]))
    elif args.edits:
        edits = json.load(open(args.edits, 'r', encoding='utf-8'))
        apply_edits(args.docx, edits)
    elif args.single:
        edit = json.loads(args.single)
        doc = Document(args.docx)
        backup(args.docx)
        apply_single_edit(doc, edit)
        doc.save(args.docx)
        print('Done')
    else:
        print('Usage: patch.py <docx> --edits edits.json')
        print('       patch.py <docx> --replace-all "old" "new"')
        print('       patch.py --find --replace-all "old" "new"')
